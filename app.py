import streamlit as st
import requests
from bs4 import BeautifulSoup
from google import genai
from google.genai import errors as genai_errors
from streamlit_local_storage import LocalStorage
import os
import time
import urllib.parse
import io
import re
import subprocess
import shutil
import glob
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK


class QuotaEpuiseeError(Exception):
    """Levée quand Gemini renvoie un 429 (quota gratuit
    épuisé) — distincte d'un 503 passager, car réessayer ne
    sert à rien avant la réinitialisation du quota."""
    pass


class ModeleIndisponibleError(Exception):
    """Levée quand Gemini renvoie un 404 pour le modèle
    demandé (identifiant retiré/déprécié côté Google) —
    permanent comme le quota épuisé : réessayer le même
    modèle ne sert à rien, il faut basculer vers un autre."""
    pass


# ============================================================
# CONFIGURATION
# ============================================================
#
# Le modèle Gemini n'est plus fixé ici : il se choisit dans
# l'interface (menu déroulant "🤖 Modèle Gemini").


# ============================================================
# PAGE STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Générateur de Cours IFSI",
    page_icon="📚",
    layout="wide"
)


# ============================================================
# FONCTION : AUTO-SCROLL PENDANT LE TRAITEMENT
# ============================================================
#
# Streamlit ne fait pas défiler la page automatiquement quand
# du nouveau contenu apparaît. On observe les changements du
# DOM et on fait défiler vers le bas à chaque ajout, le temps
# du traitement.
# ============================================================

def activer_autoscroll():

    html_autoscroll = """
        <script>
        let suivreLeBas = true;

        function trouverConteneurDefilant(doc) {
            const selecteurs = [
                '[data-testid="stAppViewContainer"]',
                '[data-testid="stMain"]',
                'section.main',
                'div.main'
            ];

            for (const sel of selecteurs) {
                const el = doc.querySelector(sel);
                if (el && el.scrollHeight > el.clientHeight) {
                    return el;
                }
            }

            return null;
        }

        function estPresDuBas(scrollTop, scrollHeight, hauteurVisible) {
            const seuil = 120;
            return (
                scrollHeight - scrollTop - hauteurVisible < seuil
            );
        }

        function surveillerDefilementManuel(doc, conteneur) {
            const fenetre = window.parent;

            function verifierPosition() {
                if (conteneur) {
                    suivreLeBas = estPresDuBas(
                        conteneur.scrollTop,
                        conteneur.scrollHeight,
                        conteneur.clientHeight
                    );
                } else {
                    suivreLeBas = estPresDuBas(
                        fenetre.scrollY,
                        doc.body.scrollHeight,
                        fenetre.innerHeight
                    );
                }
            }

            if (conteneur) {
                conteneur.addEventListener(
                    'scroll', verifierPosition, { passive: true }
                );
            } else {
                fenetre.addEventListener(
                    'scroll', verifierPosition, { passive: true }
                );
            }
        }

        function defilerVersLeBas() {
            if (!suivreLeBas) {
                return;
            }

            const doc = window.parent.document;
            const conteneur = trouverConteneurDefilant(doc);

            if (conteneur) {
                conteneur.scrollTop = conteneur.scrollHeight;
            }

            // Repli : on fait aussi défiler la fenêtre entière,
            // au cas où le vrai conteneur défilant serait
            // ailleurs (structure Streamlit qui peut varier).
            window.parent.scrollTo(
                0,
                doc.body.scrollHeight
            );
        }

        const observateur = new MutationObserver(
            defilerVersLeBas
        );

        observateur.observe(
            window.parent.document.body,
            { childList: true, subtree: true }
        );

        surveillerDefilementManuel(
            window.parent.document,
            trouverConteneurDefilant(window.parent.document)
        );

        defilerVersLeBas();
        </script>
        """

    if hasattr(st, "iframe"):

        st.iframe(
            html_autoscroll,
            height=1
        )

    else:

        # Repli pour les versions de Streamlit antérieures à
        # l'introduction de st.iframe (ex. environnements
        # locaux avec une version plus ancienne installée).
        import streamlit.components.v1 as components

        components.html(
            html_autoscroll,
            height=1
        )


# ============================================================
# FONCTION : TÉLÉCHARGER LE MP3
# ============================================================

# ============================================================
# FONCTION : TROUVER LE VRAI DOMAINE (og:url)
# ============================================================

def obtenir_base_canonique(soup, url_page):

    # Certaines pages (ex : mediacenter.univ-lyon1.fr) embarquent
    # le lecteur Nudgis/MyVideo mais ne sont pas elles-mêmes le
    # serveur qui sert les liens /api/v2/... ou /downloads/...
    # La balise <meta property="og:url"> contient l'URL canonique
    # réelle (ex : myvideo.univ-lyon1.fr) — on l'utilise comme
    # base pour résoudre les liens relatifs si elle existe.

    meta = soup.find(
        "meta",
        attrs={"property": "og:url"}
    )

    if meta and meta.get("content"):

        parsed = urllib.parse.urlparse(
            meta["content"]
        )

        if parsed.scheme and parsed.netloc:

            base = f"{parsed.scheme}://{parsed.netloc}/"

            return base

    return url_page


# ============================================================
# FONCTION : DÉTECTER / EXTRAIRE LES LIENS VIDÉO D'UNE PAGE
# ============================================================
#
# Permet à l'utilisateur de coller l'URL d'une page de cours
# (Moodle "Livre", chapitre, etc.) plutôt que de devoir
# récupérer chaque lien vidéo à la main. On distingue :
# - un lien qui pointe déjà directement vers une vidéo
#   (mediacenter/myvideo) → traité tel quel ;
# - toute autre page → on la scanne pour en extraire tous
#   les liens vidéo qu'elle contient.
# ============================================================

DOMAINES_VIDEO = [
    "mediacenter.univ-lyon1.fr",
    "myvideo.univ-lyon1.fr",
]


def nettoyer_nom_fichier(titre):

    if not titre:

        return "Cours_IFSI"

    nom = re.sub(r'[\\/:*?"<>|]', "", titre)

    nom = re.sub(r"\s+", " ", nom).strip()

    nom = nom.replace(" - ", " - ")

    if len(nom) > 120:

        nom = nom[:120].rstrip()

    return nom or "Cours_IFSI"


def titre_depuis_url_fichier(url):

    # Utilisé en dernier recours quand l'utilisateur colle des
    # liens directs (vidéo/PDF) plutôt qu'une page de cours à
    # scanner — dans ce cas, aucun <h1>/<h2> de page n'est
    # jamais lu, donc titre_cours resterait vide sans ce repli.

    try:

        chemin = urllib.parse.urlparse(url).path

        nom_fichier = urllib.parse.unquote(
            chemin.rsplit("/", 1)[-1]
        )

    except Exception:

        return None

    nom_sans_extension = re.sub(
        r"\.[a-zA-Z0-9]+$",
        "",
        nom_fichier
    )

    titre = nom_sans_extension.replace("_", " ").strip()

    return titre or None


def est_lien_video_direct(url):

    try:

        analyse = urllib.parse.urlparse(url)

    except Exception:

        return False

    if not any(
        domaine in analyse.netloc
        for domaine in DOMAINES_VIDEO
    ):

        return False

    if "video=" in analyse.query:

        return True

    if "/permalink/" in analyse.path:

        return True

    return False


def extraire_titre_page(soup):

    h1 = soup.find("h1")

    titre_h1 = (
        h1.get_text(strip=True)
        if h1
        else None
    )


    conteneur_chapitre = soup.find(
        "div",
        id="mod_book-chapter"
    )

    titre_h2 = None

    if conteneur_chapitre:

        h2 = conteneur_chapitre.find("h2")

        if h2:

            titre_h2 = h2.get_text(strip=True)


    parties = [
        t for t in (titre_h1, titre_h2)
        if t
    ]

    if not parties:

        return None

    return " - ".join(parties)


def est_lien_pdf(url):

    try:

        analyse = urllib.parse.urlparse(url)

    except Exception:

        return False

    return analyse.path.lower().endswith(".pdf")


def extraire_liens_videos_page(url_page, session):

    try:

        response = session.get(
            url_page,
            timeout=30
        )

        response.raise_for_status()

    except Exception as e:

        st.error(
            f"❌ Impossible d'accéder à la page : {e}"
        )

        return [], [], None


    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )


    liens_videos = []
    liens_pdf = []
    deja_vus = set()

    for lien in soup.find_all(
        "a",
        href=True
    ):

        href_absolu = urllib.parse.urljoin(
            url_page,
            lien["href"]
        )

        if href_absolu in deja_vus:

            continue


        if est_lien_video_direct(href_absolu):

            deja_vus.add(href_absolu)

            liens_videos.append(href_absolu)

        elif est_lien_pdf(href_absolu):

            deja_vus.add(href_absolu)

            liens_pdf.append(href_absolu)


    titre = extraire_titre_page(soup)

    return liens_videos, liens_pdf, titre


def developper_urls(urls, session):

    urls_finales = []
    pdfs_finaux = []
    titre_cours = None

    for url in urls:

        if est_lien_video_direct(url):

            urls_finales.append(url)

            continue

        if est_lien_pdf(url):

            st.write(
                f"📄 Support PDF détecté directement : {url}"
            )

            pdfs_finaux.append(url)

            continue


        st.write(
            f"🔎 Page de cours détectée, recherche des "
            f"vidéos sur : {url}"
        )

        liens_videos, liens_pdf, titre = extraire_liens_videos_page(
            url,
            session
        )

        if titre_cours is None and titre:

            titre_cours = titre


        if liens_videos:

            st.write(
                f"✅ {len(liens_videos)} vidéo(s) trouvée(s) "
                f"sur cette page."
            )

            urls_finales.extend(liens_videos)

        else:

            st.warning(
                f"⚠️ Aucun lien vidéo trouvé sur : {url}"
            )


        if liens_pdf:

            st.write(
                f"📄 {len(liens_pdf)} support(s) PDF trouvé(s) "
                f"sur cette page."
            )

            pdfs_finaux.extend(liens_pdf)


    if titre_cours is None:

        for url_pdf in pdfs_finaux:

            titre_cours = titre_depuis_url_fichier(url_pdf)

            if titre_cours:

                break

        if titre_cours is None:

            for url_video in urls_finales:

                titre_cours = titre_depuis_url_fichier(
                    url_video
                )

                if titre_cours:

                    break


    return urls_finales, pdfs_finaux, titre_cours


def telecharger_mp3(url_page, index, session):

    try:

        response = session.get(
            url_page,
            timeout=30
        )

        response.raise_for_status()

    except Exception as e:

        st.error(
            f"❌ Impossible d'accéder à la page : {e}"
        )

        return None


    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )


    # Recherche du MP3
    lien_mp3 = None

    for lien in soup.find_all(
        "a",
        href=True
    ):

        href = lien["href"]

        if ".mp3" in href.lower():

            lien_mp3 = href

            break


    if not lien_mp3:

        return None


    # Transforme une URL relative en URL absolue, en se basant
    # sur le vrai domaine (og:url) plutôt que sur l'URL collée
    # par l'utilisateur.
    base = obtenir_base_canonique(soup, url_page)

    url_mp3 = urllib.parse.urljoin(
        base,
        lien_mp3
    )


    st.write(
        "🔗 Fichier MP3 trouvé"
    )


    try:

        audio_response = session.get(
            url_mp3,
            timeout=120
        )

        audio_response.raise_for_status()

    except Exception as e:

        st.error(
            f"❌ Erreur pendant le téléchargement du MP3 : {e}"
        )

        return None


    if not audio_response.content:

        st.error(
            "❌ Le fichier MP3 est vide."
        )

        return None


    # ------------------------------------------------------
    # VÉRIFICATION : est-ce vraiment un fichier audio ?
    # ------------------------------------------------------
    # Si la session n'est pas authentifiée, le serveur peut
    # répondre 200 OK mais avec une page HTML d'erreur ou de
    # connexion à la place du MP3. On vérifie donc le
    # Content-Type et la taille avant d'accepter le fichier.

    content_type = audio_response.headers.get(
        "Content-Type", ""
    ).lower()

    taille_octets = len(audio_response.content)

    if "audio" not in content_type or taille_octets < 100_000:

        return None


    nom_fichier = (
        f"cours_ifsi_{index}.mp3"
    )


    try:

        with open(
            nom_fichier,
            "wb"
        ) as fichier:

            fichier.write(
                audio_response.content
            )

    except Exception as e:

        st.error(
            f"❌ Impossible de sauvegarder le MP3 : {e}"
        )

        return None


    taille_mo = taille_octets / 1024 / 1024

    st.write(
        f"📦 Taille : {taille_mo:.2f} Mo"
    )


    return nom_fichier


# ============================================================
# FONCTION : EXTRAIRE L'AUDIO D'UN FLUX HLS AVEC FFMPEG
# ============================================================

def extraire_audio_flux(url_flux, index, session):

    if shutil.which("ffmpeg") is None:

        st.error(
            """
❌ ffmpeg n'est pas installé sur cette machine.

Sur macOS :

brew install ffmpeg
"""
        )

        return None


    nom_fichier = (
        f"cours_ifsi_{index}.mp3"
    )


    # On ne transmet des en-têtes à ffmpeg QUE si on a un
    # vrai cookie de session. Un en-tête "Cookie:" vide suffit
    # à faire échouer la requête sur certains serveurs — d'où
    # l'importance de ne pas l'envoyer quand il n'y a rien.
    cookie_header = "; ".join(
        f"{c.name}={c.value}"
        for c in session.cookies
    )

    commande = [
        "ffmpeg",
        "-y",
        "-i", url_flux,
        "-vn",
        "-acodec", "libmp3lame",
        "-ar", "44100",
        "-b:a", "128k",
        nom_fichier
    ]

    if cookie_header.strip():

        user_agent = session.headers.get(
            "User-Agent",
            ""
        )

        entetes_ffmpeg = (
            f"Cookie: {cookie_header}\r\n"
            f"User-Agent: {user_agent}\r\n"
        )

        # On insère -headers juste après "-y", avant -i.
        commande.insert(1, entetes_ffmpeg)
        commande.insert(1, "-headers")


    st.write(
        "🎬 Extraction de l'audio depuis le flux "
        "(ffmpeg)..."
    )


    try:

        resultat = subprocess.run(
            commande,
            capture_output=True,
            text=True,
            timeout=900
        )

    except subprocess.TimeoutExpired:

        st.error(
            "⏱️ ffmpeg a mis trop de temps à extraire l'audio."
        )

        return None

    except Exception as e:

        st.error(
            f"❌ Erreur lors de l'exécution de ffmpeg : {e}"
        )

        return None


    if resultat.returncode != 0:

        st.error(
            "❌ ffmpeg n'a pas réussi à extraire l'audio "
            "du flux :"
        )

        st.code(
            resultat.stderr[-3000:]
        )

        return None


    if (
        not os.path.exists(nom_fichier)
        or os.path.getsize(nom_fichier) < 100_000
    ):

        st.error(
            "❌ Le fichier audio extrait est invalide ou "
            "trop petit."
        )

        return None


    taille_mo = (
        os.path.getsize(nom_fichier)
        / 1024
        / 1024
    )

    st.write(
        f"📦 Taille (via ffmpeg) : {taille_mo:.2f} Mo"
    )


    return nom_fichier


# ============================================================
# FONCTION : TROUVER L'OID DE LA VIDÉO
# ============================================================

def obtenir_oid(html_texte):

    correspondance = re.search(
        r'mediaOID\s*:\s*["\']([^"\']+)["\']',
        html_texte
    )

    if correspondance:

        return correspondance.group(1)

    return None


# ============================================================
# FONCTION : RÉCUPÉRER L'AUDIO VIA L'API "modes"
# ============================================================
#
# Le lien HLS statique présent dans le panneau de partage est
# parfois un point d'entrée générique qui ne fonctionne pas.
# Le lecteur JS, lui, appelle en réalité l'API /api/v2/medias/
# modes/ qui renvoie un JSON avec les vraies URLs signées
# (token "st"/"e") pour chaque qualité disponible — que ce
# soit un MP4 progressif ou un flux HLS. C'est cette méthode
# que l'on reproduit ici, car elle est beaucoup plus fiable.
# ============================================================

def recuperer_audio_via_modes(url_page, index, session):

    try:

        response = session.get(
            url_page,
            timeout=30
        )

        response.raise_for_status()

    except Exception as e:

        st.error(
            f"❌ Impossible d'accéder à la page : {e}"
        )

        return None


    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )

    base = obtenir_base_canonique(
        soup,
        url_page
    )

    oid = obtenir_oid(
        response.text
    )

    if not oid:

        st.error(
            "❌ Impossible de trouver l'identifiant (oid) "
            "de la vidéo sur cette page."
        )

        return None


    url_modes = urllib.parse.urljoin(
        base,
        f"/api/v2/medias/modes/"
        f"?oid={oid}"
        f"&html5=webm_ogg_ogv_oga_mp4_m4a_mp3_m3u8"
        f"&yt=yt&embed=embed&maxheight=1080"
    )


    try:

        reponse_modes = session.get(
            url_modes,
            timeout=30
        )

        reponse_modes.raise_for_status()

        data = reponse_modes.json()

    except Exception as e:

        st.error(
            f"❌ Impossible de récupérer les informations "
            f"vidéo (API modes) : {e}"
        )

        return None


    noms_qualites = data.get("names", [])

    if not noms_qualites:

        st.error(
            "❌ Aucune qualité vidéo disponible pour "
            "cette vidéo."
        )

        return None


    # On choisit la qualité avec le débit le plus faible :
    # on n'a besoin que de l'audio, pas de la meilleure image.
    meilleure_qualite = min(
        noms_qualites,
        key=lambda q: data.get(q, {})
                          .get("resource", {})
                          .get("bitrate", float("inf"))
    )

    ressource = data.get(
        meilleure_qualite, {}
    ).get("resource")

    if not ressource or not ressource.get("url"):

        st.error(
            "❌ Aucune URL de média trouvée dans la "
            "réponse de l'API modes."
        )

        return None


    st.write(
        f"🔗 Média trouvé (qualité {meilleure_qualite})"
    )


    return extraire_audio_flux(
        ressource["url"],
        index,
        session
    )


# ============================================================
# FONCTION : RÉCUPÉRER L'AUDIO (MP3 direct, sinon API modes)
# ============================================================

def recuperer_audio(url_page, index, session):

    st.write(
        "⬇️ Recherche du fichier audio..."
    )

    fichier_local = telecharger_mp3(
        url_page,
        index,
        session
    )

    if fichier_local:

        return fichier_local


    st.write(
        "↪️ Pas de MP3 direct exploitable — tentative via "
        "l'API modes (MP4/HLS)."
    )

    return recuperer_audio_via_modes(
        url_page,
        index,
        session
    )


# ============================================================
# FONCTION : ACCÉLÉRER L'AUDIO (réduit le coût Gemini,
# facturé au nombre de secondes d'audio)
# ============================================================

FACTEUR_ACCELERATION = 1.3
ACCELERATION_ACTIVEE = False


def accelerer_audio(fichier_entree, index, facteur=FACTEUR_ACCELERATION):

    if shutil.which("ffmpeg") is None:

        st.write(
            "⚠️ ffmpeg indisponible — envoi de l'audio à "
            "vitesse normale."
        )

        return fichier_entree


    fichier_sortie = f"cours_ifsi_{index}_x{facteur}.mp3"

    commande = [
        "ffmpeg",
        "-y",
        "-i", fichier_entree,
        "-filter:a", f"atempo={facteur}",
        "-acodec", "libmp3lame",
        "-ar", "44100",
        "-b:a", "128k",
        fichier_sortie
    ]

    try:

        resultat = subprocess.run(
            commande,
            capture_output=True,
            text=True,
            timeout=300
        )

    except Exception as e:

        st.write(
            f"⚠️ Accélération audio impossible ({e}) — "
            f"envoi à vitesse normale."
        )

        return fichier_entree


    if (
        resultat.returncode != 0
        or not os.path.exists(fichier_sortie)
        or os.path.getsize(fichier_sortie) < 10_000
    ):

        st.write(
            "⚠️ Accélération audio échouée — envoi à "
            "vitesse normale."
        )

        return fichier_entree


    try:

        os.remove(fichier_entree)

    except Exception:

        pass


    st.write(
        f"⏩ Audio accéléré ×{facteur} (réduit le coût Gemini)"
    )


    return fichier_sortie


# ============================================================
# FONCTION : ANALYSER UN AUDIO
# ============================================================

# ============================================================
# FONCTION : APPELER GEMINI AVEC REPRISE AUTOMATIQUE
# ============================================================

def appeler_gemini_avec_reprise(client, model, contents, tentatives=4):

    delais = [10, 30, 60, 120]

    for essai in range(tentatives):

        try:

            return client.models.generate_content(
                model=model,
                contents=contents
            )

        except genai_errors.ClientError as e:

            message = str(e)

            code = getattr(e, "code", None)

            if code == 429 or "RESOURCE_EXHAUSTED" in message:

                raise QuotaEpuiseeError(message) from e

            if code == 404 or "NOT_FOUND" in message:

                raise ModeleIndisponibleError(message) from e

            raise

        except genai_errors.ServerError as e:

            if essai == tentatives - 1:

                raise

            delai = delais[
                min(essai, len(delais) - 1)
            ]

            st.warning(
                f"⚠️ Gemini est momentanément surchargé "
                f"(503) — nouvelle tentative dans {delai}s "
                f"({essai + 1}/{tentatives})..."
            )

            time.sleep(delai)


def afficher_erreur_quota():

    st.error(
        """
❌ Quota gratuit quotidien épuisé pour ce modèle.

Le niveau gratuit de Gemini limite le nombre de requêtes
par jour (généralement 20/jour pour ce modèle). Tu as
atteint cette limite avec tes tests d'aujourd'hui.

**Deux options :**
- Attends la réinitialisation du quota (à minuit, heure du
  Pacifique — environ 8h-9h du matin en France selon la
  saison) ;
- Utilise une clé API liée à un compte avec facturation
  activée pour continuer dès maintenant (coût très faible,
  quelques centimes par fiche).
"""
    )


def afficher_erreur_surcharge():

    st.error(
        """
❌ Gemini reste indisponible malgré plusieurs tentatives.

Ce n'est pas lié à ton compte ni à ton quota : les serveurs
Gemini du niveau gratuit sont temporairement saturés (ça
touche tous les comptes gratuits, pas seulement le tien).

**Que faire :**
- Réessaie dans quelques minutes, ça se résorbe
  généralement vite ;
- Si c'est urgent, une clé API liée à un compte avec
  facturation activée est nettement moins sujette à ce
  genre de saturation.
"""
    )


def afficher_erreur_modele_indisponible():

    st.error(
        """
❌ Ce modèle Gemini n'est plus disponible.

Google retire ou renomme parfois des modèles. Ce n'est pas
un problème temporaire — réessayer avec le même modèle ne
fonctionnera pas.

**Que faire :** utilise le bouton ci-dessous pour basculer
vers un autre modèle de la chaîne de repli.
"""
    )


# ============================================================
# FONCTION : TRANSCRIRE L'AUDIO EN LOCAL (Whisper, sans API)
# ============================================================
#
# Remplace l'envoi de l'audio brut à Gemini : on transcrit
# nous-mêmes, sur le serveur qui héberge l'appli, avec un
# modèle Whisper téléchargé une fois puis réutilisé. Aucune
# requête réseau vers Gemini pour cette étape — donc aucun
# risque de 503/quota sur la partie la plus lourde en tokens.
# Contrepartie assumée : plus lent (minutes, pas secondes),
# et tourne sur CPU sur l'hébergement gratuit.
# ============================================================

MODELE_WHISPER_TAILLE = "base"

_modele_whisper_charge = None


def obtenir_modele_whisper():

    global _modele_whisper_charge

    if _modele_whisper_charge is None:

        from faster_whisper import WhisperModel

        st.write(
            f"📥 Chargement du modèle Whisper "
            f"« {MODELE_WHISPER_TAILLE} » (une seule fois, "
            f"peut prendre un moment la première fois)..."
        )

        _modele_whisper_charge = WhisperModel(
            MODELE_WHISPER_TAILLE,
            device="cpu",
            compute_type="int8"
        )

    return _modele_whisper_charge


def transcrire_audio_local(fichier_local, numero):

    st.write(
        f"🎙️ Transcription locale de l'audio {numero} "
        f"(Whisper, sans appel API — peut prendre plusieurs "
        f"minutes)..."
    )

    try:

        modele = obtenir_modele_whisper()

        segments, info = modele.transcribe(
            fichier_local,
            language="fr",
            vad_filter=True
        )

        texte = " ".join(
            segment.text.strip() for segment in segments
        ).strip()

    except Exception as e:

        st.error(
            f"❌ Erreur de transcription locale : {e}"
        )

        return None


    if not texte:

        st.warning(
            f"⚠️ Transcription vide pour l'audio {numero} — "
            f"fichier probablement silencieux ou illisible."
        )

        return None


    st.success(
        f"✅ Audio {numero} transcrit localement "
        f"({len(texte.split())} mots)."
    )

    return texte


# ============================================================
# FONCTION : TRANSCRIRE L'AUDIO VIA GROQ (Whisper hébergé,
# gratuit, très rapide — ~15s pour 1h d'audio)
# ============================================================
#
# Contrairement à la transcription locale, celle-ci tourne sur
# les serveurs de Groq, pas sur cet hébergement — donc aucun
# risque de bridage CPU côté Streamlit Cloud. Reste gratuite
# (niveau gratuit Groq : 20 req/min, 2000 req/jour, 28800s
# d'audio/jour). Limite Groq : 25 Mo par fichier envoyé, donc
# on découpe les gros fichiers avant envoi.
# ============================================================

LIMITE_TAILLE_GROQ_MO = 24  # marge de sécurité sous les 25 Mo


def diviser_audio_en_morceaux(fichier_local, numero, duree_segment=600):

    if shutil.which("ffmpeg") is None:

        return [fichier_local]

    prefixe = f"groq_chunk_{numero}_%03d.mp3"

    commande = [
        "ffmpeg", "-y",
        "-i", fichier_local,
        "-f", "segment",
        "-segment_time", str(duree_segment),
        "-c", "copy",
        prefixe
    ]

    try:

        subprocess.run(
            commande,
            check=True,
            capture_output=True,
            timeout=120
        )

    except Exception as e:

        st.write(
            f"⚠️ Échec du découpage audio : {e} — envoi du "
            f"fichier entier tel quel (risque d'échec Groq "
            f"si > 25 Mo)."
        )

        return [fichier_local]

    morceaux = sorted(
        glob.glob(f"groq_chunk_{numero}_*.mp3")
    )

    return morceaux if morceaux else [fichier_local]


def transcrire_audio_groq(fichier_local, numero, cle_api_groq):

    try:

        from groq import Groq

        client_groq = Groq(api_key=cle_api_groq)

    except Exception as e:

        st.write(
            f"⚠️ Impossible d'initialiser le client Groq : {e}"
        )

        return None


    taille_mo = os.path.getsize(fichier_local) / (1024 * 1024)

    if taille_mo <= LIMITE_TAILLE_GROQ_MO:

        morceaux = [fichier_local]

    else:

        st.write(
            f"✂️ Audio {numero} trop volumineux pour Groq "
            f"({taille_mo:.1f} Mo) — découpage en morceaux..."
        )

        morceaux = diviser_audio_en_morceaux(
            fichier_local,
            numero
        )


    st.write(
        f"🎙️ Transcription de l'audio {numero} via Groq "
        f"(Whisper hébergé, quelques secondes)..."
    )

    textes = []

    for idx, morceau in enumerate(morceaux):

        try:

            with open(morceau, "rb") as f:

                transcription = (
                    client_groq.audio.transcriptions.create(
                        file=(os.path.basename(morceau), f.read()),
                        model="whisper-large-v3",
                        language="fr",
                        response_format="text"
                    )
                )

            textes.append(str(transcription).strip())

        except Exception as e:

            st.write(
                f"⚠️ Erreur Groq sur le morceau {idx + 1} de "
                f"l'audio {numero} : {e}"
            )

            for m in morceaux:

                if m != fichier_local:

                    try:

                        os.remove(m)

                    except Exception:

                        pass

            return None

        finally:

            if morceau != fichier_local:

                try:

                    os.remove(morceau)

                except Exception:

                    pass


    texte_final = " ".join(
        t for t in textes if t
    ).strip()

    if not texte_final:

        st.warning(
            f"⚠️ Transcription Groq vide pour l'audio {numero}."
        )

        return None


    st.success(
        f"✅ Audio {numero} transcrit via Groq "
        f"({len(texte_final.split())} mots)."
    )

    return texte_final


def transcrire_audio(fichier_local, numero, cle_api_groq):

    # Groq en priorité (rapide, gratuit, hors de notre
    # hébergement) — repli automatique sur Whisper local si
    # pas de clé fournie ou si Groq échoue pour une raison
    # quelconque (quota, panne, fichier trop long malgré le
    # découpage...).

    if cle_api_groq:

        texte = transcrire_audio_groq(
            fichier_local,
            numero,
            cle_api_groq
        )

        if texte:

            return texte

        st.write(
            f"↪️ Groq indisponible pour l'audio {numero} — "
            f"repli sur la transcription locale (plus lente)."
        )

    return transcrire_audio_local(fichier_local, numero)


# ============================================================
# FONCTION : TÉLÉCHARGER ET EXTRAIRE LE TEXTE D'UN PDF SUPPORT
# ============================================================
#
# Les diapositives (PDF "Support" liés sur la page de cours)
# contiennent parfois des informations que le professeur n'a
# pas développées à l'oral. On extrait le texte nous-mêmes
# (bibliothèque Python, pas d'IA) plutôt que d'envoyer le PDF
# à Gemini — aucun appel réseau vers Gemini pour cette étape.
#
# Limite assumée : ne « voit » pas les schémas/images des
# diapositives, seulement le texte qu'elles contiennent.
# ============================================================

def telecharger_et_extraire_pdf(
    url_pdf,
    index,
    session,
    retourner_bytes=False
):

    # retourner_bytes=False (défaut) : comportement d'origine,
    # ne renvoie que le texte extrait (utilisé par le pipeline
    # normal mode gratuit).
    # retourner_bytes=True : renvoie le couple (texte, octets
    # bruts du PDF) — utilisé par le paquet de téléchargement
    # "sans IA", qui doit inclure le PDF original tel quel, pas
    # seulement son texte extrait.

    def echec():

        return (None, None) if retourner_bytes else None


    try:

        reponse = session.get(
            url_pdf,
            timeout=60
        )

        reponse.raise_for_status()

    except Exception as e:

        st.write(
            f"⚠️ Support PDF {index} inaccessible : {e}"
        )

        return echec()


    content_type = reponse.headers.get(
        "Content-Type", ""
    ).lower()

    if (
        "pdf" not in content_type
        and not reponse.content[:4] == b"%PDF"
    ):

        st.write(
            f"⚠️ Support {index} ne semble pas être un PDF "
            f"valide — ignoré."
        )

        return echec()


    st.write(
        f"📄 Extraction du texte du support PDF {index}..."
    )

    try:

        import pdfplumber

        with pdfplumber.open(
            io.BytesIO(reponse.content)
        ) as pdf:

            pages_texte = [
                page.extract_text() or ""
                for page in pdf.pages
            ]

        texte = "\n\n".join(
            t for t in pages_texte if t.strip()
        ).strip()

    except Exception as e:

        st.write(
            f"⚠️ Erreur d'extraction du support {index} : {e}"
        )

        return echec()


    if not texte:

        st.warning(
            f"⚠️ Aucun texte extrait du support {index} "
            f"(PDF probablement composé uniquement d'images "
            f"scannées, non détectable par extraction directe)."
        )

        # Le texte est vide mais le PDF lui-même reste valide
        # et utile dans le paquet de téléchargement — on le
        # renvoie quand même dans ce cas.
        return (None, reponse.content) if retourner_bytes else None


    st.success(
        f"✅ Support PDF {index} extrait "
        f"({len(texte.split())} mots)."
    )

    return (texte, reponse.content) if retourner_bytes else texte


# ============================================================
# FONCTIONS DU MODE "RAPIDE" : envoi direct à Gemini (audio
# et PDF bruts), sans transcription/extraction locale. Plus
# rapide, mais consomme davantage de quota Gemini par cours.
# ============================================================

def attendre_fichier_gemini(client, fichier):

    maximum = 300

    temps = 0

    while temps < maximum:

        try:

            info = client.files.get(
                name=fichier.name
            )

        except Exception as e:

            st.error(
                f"❌ Impossible de vérifier le fichier : {e}"
            )

            return None


        if hasattr(info.state, "name"):

            statut = info.state.name

        else:

            statut = str(info.state)


        if statut == "ACTIVE":

            # On retourne l'objet ORIGINAL renvoyé par
            # client.files.upload() (celui reçu en paramètre),
            # pas "info" (celui de client.files.get()) — le
            # passer à generate_content() déclenche un 400.
            return fichier


        if statut in ["FAILED", "ERROR"]:

            st.error(
                f"❌ Gemini n'a pas réussi à traiter "
                f"{fichier.name}"
            )

            return None


        time.sleep(3)

        temps += 3


    st.error(
        "⏱️ Gemini a mis trop de temps à traiter le fichier."
    )

    return None


def uploader_audio_gemini(client, fichier_local, numero):

    st.write(
        f"⬆️ Envoi de l'audio {numero} à Gemini..."
    )

    try:

        fichier = client.files.upload(
            file=fichier_local
        )

    except Exception as e:

        st.error(
            f"❌ Erreur d'upload : {e}"
        )

        return None


    fichier_pret = attendre_fichier_gemini(
        client,
        fichier
    )

    if fichier_pret is None:

        try:

            client.files.delete(name=fichier.name)

        except Exception:

            pass

        return None


    st.success(
        f"✅ Audio {numero} envoyé et prêt."
    )

    return fichier_pret


def telecharger_et_uploader_pdf(
    url_pdf,
    index,
    session,
    client,
    retourner_bytes=False
):

    def echec():

        return (None, None) if retourner_bytes else None


    try:

        reponse = session.get(
            url_pdf,
            timeout=60
        )

        reponse.raise_for_status()

    except Exception as e:

        st.write(
            f"⚠️ Support PDF {index} inaccessible : {e}"
        )

        return echec()


    content_type = reponse.headers.get(
        "Content-Type", ""
    ).lower()

    if (
        "pdf" not in content_type
        and not reponse.content[:4] == b"%PDF"
    ):

        st.write(
            f"⚠️ Support {index} ne semble pas être un PDF "
            f"valide — ignoré."
        )

        return echec()


    nom_fichier = f"support_{index}.pdf"

    try:

        with open(nom_fichier, "wb") as f:

            f.write(reponse.content)

    except Exception as e:

        st.write(
            f"⚠️ Impossible d'enregistrer le support {index} : {e}"
        )

        return echec()


    st.write(
        f"⬆️ Envoi du support PDF {index} à Gemini..."
    )

    try:

        fichier = client.files.upload(
            file=nom_fichier
        )

    except Exception as e:

        st.write(
            f"⚠️ Erreur d'upload du support {index} : {e}"
        )

        return echec()

    finally:

        try:

            os.remove(nom_fichier)

        except Exception:

            pass


    fichier_pret = attendre_fichier_gemini(
        client,
        fichier
    )

    if fichier_pret is None:

        try:

            client.files.delete(name=fichier.name)

        except Exception:

            pass

        return echec()


    st.success(
        f"✅ Support PDF {index} envoyé et prêt."
    )

    if retourner_bytes:

        return fichier_pret, reponse.content

    return fichier_pret


# ============================================================
# FONCTION : CRÉER LA FICHE FINALE (un seul appel Gemini,
# texte seul — transcriptions + PDF déjà obtenus en local)
# ============================================================
#
# L'audio est transcrit localement (Whisper, sans appel API)
# et les PDF sont extraits localement (pdfplumber) avant même
# d'arriver ici. Gemini ne reçoit donc que du texte, jamais de
# fichier audio/PDF brut — ce qui élimine toute exposition aux
# 503/quota sur la partie la plus lourde du traitement, et
# réduit fortement le volume de tokens envoyé (le texte est
# beaucoup plus compact que l'audio brut).
#
# Compromis assumé : si cet appel échoue, tout est à refaire
# pour la fiche (mais pas les transcriptions, conservées pour
# permettre une reprise avec un autre modèle).
# ============================================================

# ============================================================
# STRUCTURE ET RÈGLES DE LA FICHE FINALE (partagée entre
# l'appel direct et la fusion depuis des notes de sous-lots)
# ============================================================

STRUCTURE_ET_REGLES_FICHE = r"""
==============================
POSTURE SOIGNANTE — TRÈS IMPORTANT
==============================

Cette fiche est destinée à un(e) étudiant(e) INFIRMIER(ÈRE),
pas à un étudiant en médecine ou un externe. Ne te contente
pas de résumer la théorie médicale : relie systématiquement
la physiopathologie à la clinique infirmière — ce que ça
change concrètement pour la surveillance, les gestes et les
priorités d'un(e) IDE.

Imagine aussi que tu es cet(te) étudiant(e) en train de
prendre des notes PENDANT le cours magistral lui-même — pas
en train de réécrire un manuel après coup. Ça implique :

- Ne développe PAS en détail les mécanismes moléculaires ou
  immunologiques fins (cascades biochimiques, sous-types de
  cellules, voies de signalisation...) sauf s'ils ont une
  conséquence clinique directe et utile à l'exercice
  infirmier. Une phrase de mécanisme général suffit la
  plupart du temps ("lésion de l'endothélium → dépôt de
  fibrine → greffe bactérienne" plutôt qu'un paragraphe sur
  la cascade du complément).
- Privilégie systématiquement ce qui est ACTIONNABLE pour un
  infirmier : que surveiller, quand alerter, quel geste
  technique, quelle précaution, quelle règle clinique
  simple à retenir.
- Capture les INSISTANCES ORALES de l'enseignant quand elles
  sont perceptibles dans l'audio (un point répété, un
  avertissement du type "ça, c'est un piège classique à
  l'examen", un seuil qu'il souligne) — c'est exactement ce
  qu'un(e) bon(ne) étudiant(e) noterait en marge pendant le
  cours. Ne vise pas l'exhaustivité à 100% du support, mais
  ne loupe pas ce que l'enseignant signale comme important.
- Utilise un vocabulaire clair, évite le jargon médical non
  expliqué.

==============================
RIGUEUR BIOLOGIQUE
==============================

- Quand une valeur biologique est mentionnée (numération,
  ionogramme, etc.), donne-la avec son UNITÉ EXACTE telle
  qu'énoncée dans le cours (ex. : G/L et non g/L pour une
  numération leucocytaire, mmol/L et non µmol/L) — ne
  généralise pas et ne corrige pas silencieusement une
  unité, mais reste fidèle à ce qui est dit.
- Signale explicitement les seuils d'alerte vitale évoqués
  dans le cours (ex. : agranulocytose, hyperkaliémie,
  thrombopénie sévère...) avec leur valeur si elle est
  donnée.

==============================
STRUCTURE OBLIGATOIRE
==============================

# 1. Physiopathologie (l'essentiel)

Pour chaque pathologie ou notion abordée dans le cours,
donne en quelques lignes maximum :

- le mécanisme en une phrase simple ;
- les conséquences principales ;
- les éléments clés (localisation, germes en cause,
  facteurs favorisants...) si pertinents.

Reste concis — c'est un rappel, pas un cours d'anatomo-
pathologie. Utilise des sous-titres par pathologie/notion.

Si l'enseignant signale un piège d'examen classique ou une
confusion terminologique fréquente à propos d'une notion de
cette section, ajoute juste après un mini-titre "### Piège
classique au partiel IFSI" suivi d'une ou deux puces
expliquant le piège — uniquement quand c'est réellement
mentionné ou clairement déductible du cours, pas
systématiquement pour chaque notion.


# 2. Signes cliniques et d'alerte

Présente les pathologies du cours sous forme de TABLEAU
avec exactement ces colonnes :

| Pathologie | Signes cliniques (au lit du patient) | Signes paracliniques (examens) | Signes de gravité / Complications |

- Colonne "Signes cliniques" : uniquement ce qu'un(e)
  infirmier(ère) observe ou recueille directement auprès du
  patient (inspection, palpation, interrogatoire,
  constantes) — pas de résultat d'examen.
- Colonne "Signes paracliniques" : résultats de biologie,
  imagerie ou autres examens complémentaires, avec unités
  exactes quand elles sont données dans le cours. Si le
  cours ne mentionne aucun signe paraclinique pour une
  pathologie, laisse la cellule avec un tiret "—".
- Une ligne par pathologie ou situation clinique abordée
  dans le cours.


# 3. Rôle propre et surveillance infirmière (IDE)

Organise cette section en sous-parties claires, par exemple :

- Surveillance des paramètres vitaux (lesquels, pourquoi,
  quel seuil d'alerte) ;
- Prélèvements et examens (quoi, quand, précautions avant/
  après le geste) ;
- Isolement et prévention (type de précautions — gouttelettes,
  contact, air — et pour quelles pathologies du cours) ;
- Tout autre élément de rôle propre pertinent évoqué dans
  le cours (soins, gestes techniques, éducation du patient).

Reste concret et actionnable : ce qu'un(e) infirmier(ère)
fait réellement, pas de la théorie médicale.


# 4. Urgences absolues et règles d'or

Termine par deux sous-parties :

- "Urgences vitales absolues" : les situations décrites
  dans le cours qui nécessitent une action immédiate
  (avec le geste ou le traitement d'urgence associé s'il
  est mentionné) ;
- "Règles d'or" : des phrases courtes et mémorisables du
  type "Toute fièvre chez X = Y jusqu'à preuve du
  contraire", reprenant les raccourcis cliniques donnés
  par l'enseignant ou déductibles clairement du cours.

N'invente aucune urgence ou règle qui ne serait pas
mentionnée ou clairement déductible du contenu fourni.


==============================
RÈGLES
==============================

- Reste fidèle au contenu fourni, n'invente aucune
  information.
- Fusionne les informations et supprime les répétitions.
- Ne perds aucune information importante, y compris les
  listes ou exemples présents uniquement sur une diapositive.
- Si deux passages semblent contradictoires,
  signale-le plutôt que d'inventer une réponse.
- Si un passage est incompréhensible, indique-le brièvement.
- Utilise des tableaux Markdown quand ils facilitent la
  lecture (obligatoire pour la section 2, avec exactement
  les 4 colonnes demandées).
- Utilise des listes à puces courtes plutôt que des
  paragraphes denses.
- Mets en gras (**...**) les termes et valeurs clés
  (seuils, noms de pathologies, règles d'or) pour que la
  fiche soit rapide à scanner visuellement.
- N'utilise JAMAIS de notation LaTeX (pas de \text{},
  \circ, \,, ^{...}, _{...}, $...$, etc.). Écris tout en
  texte normal lisible : "38,5 °C" (pas "38,5\,^\circ\text{C}"),
  "CD4+" (pas "CD4^{+}"), "10 puissance 14" ou "10^14" en
  toutes lettres (pas "10^{14}"), "SpO2" (pas "\text{SpO}_2").
- La fiche doit rester dense et actionnable, pas exhaustive
  au sens "cours magistral" — un étudiant doit pouvoir la
  relire juste avant un examen ou un stage.
"""


# Nombre maximal d'éléments (transcriptions audio + textes
# PDF confondus) envoyés dans un seul appel Gemini. Au-delà,
# on découpe en sous-lots pour limiter l'exposition aux 503
# sur les grosses requêtes.
#
# Remonté de 8 à 10 : plus le nombre d'appels Gemini
# séquentiels nécessaires pour une fiche est élevé, plus la
# probabilité cumulée de tomber sur un 503 "high demand"
# augmente (chaque appel est une chance en plus d'échouer),
# même quand aucun plafond de quota n'est réellement dépassé.
#
# ⚠️ Recalibrage à prévoir : les chiffres ci-dessous
# (~18K-35K tokens/élément) dataient de l'ancienne
# architecture, où l'audio brut était envoyé à Gemini. Depuis
# le passage à la transcription locale (Whisper) + extraction
# PDF locale, chaque élément n'est plus que du TEXTE — bien
# plus compact que l'audio/PDF brut. Le vrai plafond de
# tokens/minute est donc probablement beaucoup plus loin
# qu'avant ; ce seuil de 10 reste hérité de l'ancienne
# architecture et mériterait d'être retesté (voire remonté
# nettement) une fois des données réelles disponibles sur le
# nouveau pipeline.
#
# Nos observations (ancienne architecture) donnaient une
# fourchette approximative de ~18K à ~35K tokens/fichier
# selon la durée des audios — donc ce seuil reste un
# compromis empirique, pas une garantie absolue de rester
# sous les 250K tokens/minute du niveau gratuit.
#
# Note : un test à 999 (découpage désactivé) a été tenté sur
# un lot de 15 fichiers pour départager "1 gros appel" vs
# "plusieurs appels séquentiels", mais il a buté sur un 429
# (quota journalier de 20 requêtes/jour épuisé, cumulé par
# tous les tests de la journée) avant de pouvoir observer un
# éventuel dépassement de tokens — donc non concluant sur ce
# point précis. Remis à 10 par prudence.
TAILLE_MAX_SOUS_LOT = 10


def construire_prompt_fiche_texte(contenus_textuels):

    # Isolé de creer_fiche_finale_texte() pour pouvoir réutiliser
    # exactement le même prompt (consignes + contenu) dans le
    # bouton de téléchargement de secours — l'utilisateur doit
    # pouvoir coller tel quel dans une autre IA, sans reconstruire
    # les instructions lui-même.

    contenu_source = "\n\n".join(contenus_textuels)

    return f"""
Tu es un formateur expert en Institut de Formation
en Soins Infirmiers (IFSI).

Un enseignant (souvent médecin ou expert du domaine) a
donné ce cours à des étudiants INFIRMIERS. Ta mission est
d'adapter ce contenu d'expert en fiche de révision de
niveau infirmier — pas de le retranscrire tel quel.

Tu vas recevoir la TRANSCRIPTION TEXTUELLE d'un ou plusieurs
enregistrements audio du même cours (obtenue par
reconnaissance vocale automatique — elle peut donc contenir
de légères erreurs de transcription, notamment sur des
termes médicaux techniques ; utilise ton jugement clinique
pour corriger silencieusement les erreurs évidentes de ce
type, ex. un terme médical mal orthographié phonétiquement).

Si des PDF de support ("diapositives") existent pour ce
cours, ils te sont fournis SÉPARÉMENT (en fichiers joints
directement à cette requête, pas dans le texte ci-dessous) —
lis-les attentivement, schémas et images compris.

IMPORTANT : certaines informations (listes, exemples,
pathologies citées) peuvent apparaître UNIQUEMENT dans un
support PDF sans avoir été développées à l'oral, ou
inversement UNIQUEMENT dans la transcription orale sans
être écrites sur une diapositive. Croise systématiquement
les deux sources et n'omets aucune liste ou exemple clinique
présent dans l'une ou l'autre, même brièvement.

==============================
CONTENU DU COURS
==============================

{contenu_source}
""" + STRUCTURE_ET_REGLES_FICHE


def creer_fiche_finale_texte(
    client,
    contenus_textuels,
    model_name,
    fichiers_pdf_gemini=None
):

    st.write(
        "### 🧠 Analyse des transcriptions et création de "
        "la fiche de révision"
    )


    prompt = construire_prompt_fiche_texte(contenus_textuels)

    # Les PDF de support sont fournis à Gemini en fichiers
    # natifs (schémas/images inclus), même en mode gratuit —
    # seul l'audio est transcrit en local/Groq. L'upload en
    # lui-même ne consomme aucun quota de generate_content :
    # ça reste un seul appel, juste avec ces fichiers en plus.
    if fichiers_pdf_gemini:

        contenu_requete = [prompt] + list(fichiers_pdf_gemini)

    else:

        contenu_requete = prompt


    try:

        reponse = appeler_gemini_avec_reprise(
            client,
            model_name,
            contenu_requete
        )

    except QuotaEpuiseeError:

        raise

    except ModeleIndisponibleError:

        raise

    except genai_errors.ServerError:

        raise

    except Exception as e:

        st.error(
            "❌ Erreur Gemini lors de la création "
            "de la fiche finale."
        )

        st.exception(e)

        return None


    try:

        return reponse.text

    except Exception:

        return None


# ============================================================
# FONCTION : CRÉER LA FICHE FINALE — MODE RAPIDE (un seul
# appel Gemini, avec tous les fichiers audio/PDF directement
# en entrée, sans transcription/extraction locale)
# ============================================================

def creer_fiche_finale_fichiers(
    client,
    fichiers_geminis,
    model_name,
    fichiers_pdf_gemini=None
):

    st.write(
        "### 🧠 Analyse des audios et création de la "
        "fiche de révision"
    )


    prompt = """
Tu es un formateur expert en Institut de Formation
en Soins Infirmiers (IFSI).

Un enseignant (souvent médecin ou expert du domaine) a
donné ce cours à des étudiants INFIRMIERS. Ta mission est
d'adapter ce contenu d'expert en fiche de révision de
niveau infirmier — pas de le retranscrire tel quel.

Tu vas recevoir un ou plusieurs enregistrements audio
appartenant au même cours (éventuellement en plusieurs
parties), et éventuellement les PDF des diapositives
("supports") utilisées pendant ce cours.

Écoute les audios ET lis attentivement les PDF fournis —
ne produis PAS de retranscription intermédiaire, rédige
directement la fiche de révision finale.

IMPORTANT : certaines informations (listes, exemples,
pathologies citées) peuvent apparaître UNIQUEMENT sur une
diapositive du PDF sans avoir été développées à l'oral, ou
inversement UNIQUEMENT à l'oral sans être écrites sur la
diapositive. Croise systématiquement les deux sources et
n'omets aucune diapositive contenant une liste ou des
exemples cliniques, même si elle n'a été que brièvement
survolée pendant le cours.
""" + STRUCTURE_ET_REGLES_FICHE


    contenu_requete = [prompt] + list(fichiers_geminis)


    try:

        reponse = appeler_gemini_avec_reprise(
            client,
            model_name,
            contenu_requete
        )

    except QuotaEpuiseeError:

        raise

    except ModeleIndisponibleError:

        raise

    except genai_errors.ServerError:

        raise

    except Exception as e:

        st.error(
            "❌ Erreur Gemini lors de la création "
            "de la fiche finale."
        )

        st.exception(e)

        return None


    try:

        return reponse.text

    except Exception:

        return None


# ============================================================
# FONCTION : NOTES DENSES D'UN SOUS-LOT (transcriptions + PDF)
# ============================================================
#
# Utilisée quand il y a trop de contenu pour un seul appel.
# Produit des notes condensées (pas la fiche finale mise en
# forme) à partir d'un sous-ensemble des transcriptions/PDF
# du cours.
# ============================================================

def creer_notes_sous_lot_texte(
    client,
    contenus_sous_lot,
    model_name,
    numero_lot,
    total_lots
):

    st.write(
        f"### 🧠 Analyse du lot {numero_lot}/{total_lots} "
        f"({len(contenus_sous_lot)} élément(s))"
    )


    contenu_source = "\n\n".join(contenus_sous_lot)

    prompt = rf"""
Tu es un formateur expert en Institut de Formation
en Soins Infirmiers (IFSI).

Un enseignant (souvent médecin ou expert du domaine) donne
ce cours à des étudiants INFIRMIERS, pas à des étudiants en
médecine. Garde ça à l'esprit même à ce stade de prise de
notes : ne recopie pas mécaniquement chaque détail
moléculaire ou médical fin de l'enseignant si ça n'a pas
d'utilité clinique pour un(e) infirmier(ère) — mais
n'omets aucune pathologie, liste ou exemple clinique.

Tu vas recevoir un SOUS-ENSEMBLE (lot {numero_lot}/{total_lots})
des transcriptions audio et/ou textes de supports PDF d'un
même cours. Les transcriptions viennent d'une reconnaissance
vocale automatique — elles peuvent contenir de légères
erreurs, notamment sur des termes médicaux ; corrige
silencieusement les erreurs phonétiques évidentes. D'autres
lots de ce même cours seront traités séparément puis
fusionnés avec celui-ci pour produire la fiche de révision
finale — ce n'est PAS ton rôle ici de produire cette fiche
finale.

Lis attentivement tout le contenu fourni dans CE LOT
UNIQUEMENT.

Pour chaque pathologie ou notion abordée dans ce lot, note
de façon DENSE et STRUCTURÉE (puces courtes, style
télégraphique, pas de phrases longues ni de mise en forme
finale) :

- Mécanisme clé (physiopathologie) en une ligne ;
- Signes cliniques typiques et signes de gravité ;
- Surveillance infirmière (IDE) pertinente évoquée ;
- Urgences ou règles cliniques mentionnées.

RÈGLES :

- Reste fidèle au contenu fourni, n'invente rien.
- Sois exhaustif sur le CONTENU (n'omets aucune pathologie
  ni aucune liste présente sur une diapositive, même
  brièvement survolée à l'oral) mais très concis dans la
  FORMULATION.
- Si un passage est incompréhensible, indique-le brièvement.
- N'utilise JAMAIS de notation LaTeX (pas de \text, \circ,
  exposants/indices entre accolades, symboles $) — écris
  tout en texte normal ("38,5 °C", "CD4+", "SpO2"...).
- Ne produis pas encore de tableau ni de section "règles
  d'or" mises en forme — ce sera fait lors de la fusion
  finale avec les autres lots.

==============================
CONTENU DU LOT
==============================

{contenu_source}
"""


    try:

        reponse = appeler_gemini_avec_reprise(
            client,
            model_name,
            prompt
        )

    except QuotaEpuiseeError:

        raise

    except ModeleIndisponibleError:

        raise

    except genai_errors.ServerError:

        raise

    except Exception as e:

        st.error(
            f"❌ Erreur Gemini sur le lot {numero_lot} : {e}"
        )

        return None


    try:

        return reponse.text

    except Exception:

        return None


# ============================================================
# FONCTION : NOTES DENSES D'UN SOUS-LOT — MODE RAPIDE
# (audio + PDF envoyés directement à Gemini)
# ============================================================

def creer_notes_sous_lot_fichiers(
    client,
    fichiers_sous_lot,
    model_name,
    numero_lot,
    total_lots
):

    st.write(
        f"### 🧠 Analyse du lot {numero_lot}/{total_lots} "
        f"({len(fichiers_sous_lot)} fichier(s))"
    )


    prompt = rf"""
Tu es un formateur expert en Institut de Formation
en Soins Infirmiers (IFSI).

Un enseignant (souvent médecin ou expert du domaine) donne
ce cours à des étudiants INFIRMIERS, pas à des étudiants en
médecine. Garde ça à l'esprit même à ce stade de prise de
notes : ne recopie pas mécaniquement chaque détail
moléculaire ou médical fin de l'enseignant si ça n'a pas
d'utilité clinique pour un(e) infirmier(ère) — mais
n'omets aucune pathologie, liste ou exemple clinique.

Tu vas recevoir un SOUS-ENSEMBLE (lot {numero_lot}/{total_lots})
des enregistrements audio et/ou PDF de supports d'un même
cours. D'autres lots de ce même cours seront traités
séparément puis fusionnés avec celui-ci pour produire la
fiche de révision finale — ce n'est PAS ton rôle ici de
produire cette fiche finale.

Écoute les audios ET lis attentivement les PDF fournis dans
CE LOT UNIQUEMENT.

Pour chaque pathologie ou notion abordée dans ce lot, note
de façon DENSE et STRUCTURÉE (puces courtes, style
télégraphique, pas de phrases longues ni de mise en forme
finale) :

- Mécanisme clé (physiopathologie) en une ligne ;
- Signes cliniques typiques et signes de gravité ;
- Surveillance infirmière (IDE) pertinente évoquée ;
- Urgences ou règles cliniques mentionnées.

RÈGLES :

- Reste fidèle au contenu fourni, n'invente rien.
- Sois exhaustif sur le CONTENU (n'omets aucune pathologie
  ni aucune liste présente sur une diapositive, même
  brièvement survolée à l'oral) mais très concis dans la
  FORMULATION.
- Si un passage est incompréhensible, indique-le brièvement.
- N'utilise JAMAIS de notation LaTeX (pas de \text, \circ,
  exposants/indices entre accolades, symboles $) — écris
  tout en texte normal ("38,5 °C", "CD4+", "SpO2"...).
- Ne produis pas encore de tableau ni de section "règles
  d'or" mises en forme — ce sera fait lors de la fusion
  finale avec les autres lots.
"""


    contenu_requete = [prompt] + list(fichiers_sous_lot)


    try:

        reponse = appeler_gemini_avec_reprise(
            client,
            model_name,
            contenu_requete
        )

    except QuotaEpuiseeError:

        raise

    except ModeleIndisponibleError:

        raise

    except genai_errors.ServerError:

        raise

    except Exception as e:

        st.error(
            f"❌ Erreur Gemini sur le lot {numero_lot} : {e}"
        )

        return None


    try:

        return reponse.text

    except Exception:

        return None


# ============================================================
# FONCTION : FICHE FINALE À PARTIR DE NOTES DE SOUS-LOTS
# ============================================================
#
# Appel texte uniquement (pas de fichier audio/PDF ici) : les
# notes des sous-lots ont déjà "digéré" le contenu brut, donc
# cet appel reste léger en tokens malgré le nombre de vidéos
# initial.
# ============================================================

def creer_fiche_depuis_notes(
    client,
    notes_par_lot,
    model_name,
    fichiers_pdf_gemini=None
):

    st.write(
        "### 🧠 Fusion des lots et création de la fiche "
        "de révision finale"
    )

    contenu_notes = "\n\n".join(
        f"""
==============================
NOTES DU LOT {i + 1}
==============================

{texte}
"""
        for i, texte in enumerate(notes_par_lot)
    )

    prompt = rf"""
Tu es un formateur expert en Institut de Formation
en Soins Infirmiers (IFSI).

Un enseignant (souvent médecin ou expert du domaine) a donné
un cours à des étudiants infirmiers. Ce cours a été découpé
en plusieurs lots pour l'analyse, et chaque lot a déjà été
résumé en notes denses par un premier passage — mais ces
notes restent une simple étape intermédiaire, pas la
matière première de ta réflexion.

TA MISSION : à partir de ces notes, reconstitue et ADAPTE le
contenu du cours original pour en faire une fiche de
révision de niveau INFIRMIER — pas médecin. Tu ne
"complètes" pas des notes abstraites : tu transformes le
savoir d'un expert en un outil clinique pensé pour un(e)
étudiant(e) infirmier(ère), en gardant à l'esprit tout du
long qui est le public final.

Fusionne les différents lots, supprime les répétitions, et
ne perds aucune information importante — mais n'invente
jamais une information absente des notes.

==============================
NOTES DES DIFFÉRENTS LOTS
==============================

{contenu_notes}
""" + STRUCTURE_ET_REGLES_FICHE


    if fichiers_pdf_gemini:

        contenu_requete = [prompt] + list(fichiers_pdf_gemini)

    else:

        contenu_requete = prompt


    try:

        reponse = appeler_gemini_avec_reprise(
            client,
            model_name,
            contenu_requete
        )

    except QuotaEpuiseeError:

        raise

    except ModeleIndisponibleError:

        raise

    except genai_errors.ServerError:

        raise

    except Exception as e:

        st.error(
            "❌ Erreur Gemini lors de la fusion finale."
        )

        st.exception(e)

        return None


    try:

        return reponse.text

    except Exception:

        return None


# ============================================================
# FONCTION : ORCHESTRATEUR — appel direct ou découpage
# en sous-lots selon le nombre de fichiers
# ============================================================

def generer_fiche_finale(
    client,
    contenus,
    model_name,
    mode,
    fichiers_pdf_gemini=None
):

    if mode == "rapide":

        creer_fiche_fn = creer_fiche_finale_fichiers
        creer_notes_fn = creer_notes_sous_lot_fichiers

    else:

        creer_fiche_fn = creer_fiche_finale_texte
        creer_notes_fn = creer_notes_sous_lot_texte


    # Les PDF Gemini supplémentaires ne concernent que le mode
    # gratuit (en mode rapide, les PDF sont déjà mélangés aux
    # audios dans "contenus"). Ils sont toujours attachés au
    # DERNIER appel Gemini — direct ou fusion — jamais répétés
    # dans chaque sous-lot, pour ne pas les envoyer plusieurs
    # fois inutilement.
    pdf_pour_appel_direct = (
        fichiers_pdf_gemini if mode == "gratuit" else None
    )

    if len(contenus) <= TAILLE_MAX_SOUS_LOT:

        return creer_fiche_fn(
            client,
            contenus,
            model_name,
            pdf_pour_appel_direct
        )


    st.write(
        f"↪️ {len(contenus)} éléments au total — "
        f"découpage en sous-lots de {TAILLE_MAX_SOUS_LOT} "
        f"maximum."
    )

    sous_lots = [
        contenus[i:i + TAILLE_MAX_SOUS_LOT]
        for i in range(
            0, len(contenus), TAILLE_MAX_SOUS_LOT
        )
    ]

    notes_par_lot = []

    for i, sous_lot in enumerate(sous_lots):

        notes = creer_notes_fn(
            client,
            sous_lot,
            model_name,
            i + 1,
            len(sous_lots)
        )

        if notes:

            notes_par_lot.append(notes)

        else:

            st.warning(
                f"⚠️ Le lot {i + 1}/{len(sous_lots)} n'a pas "
                f"pu être analysé — il sera absent de la "
                f"fiche finale."
            )


    if not notes_par_lot:

        st.error(
            "❌ Aucun lot n'a pu être analysé."
        )

        return None


    # La fusion depuis les notes est toujours en texte seul,
    # quel que soit le mode d'origine (audio/PDF ou
    # transcription) — les notes de sous-lots sont déjà du
    # texte dans les deux cas. Les PDF Gemini (mode gratuit),
    # eux, n'ont pas été inclus dans les sous-lots — on les
    # attache ici, à ce dernier appel.
    return creer_fiche_depuis_notes(
        client,
        notes_par_lot,
        model_name,
        pdf_pour_appel_direct
    )


# ============================================================
# FONCTION : WORD (conversion Markdown → docx)
# ============================================================

# ============================================================
# THÈME VISUEL DU DOCUMENT
# ============================================================

COULEUR_TITRE_PRINCIPAL = RGBColor(0x1F, 0x4E, 0x79)   # bleu foncé
COULEUR_SOUS_TITRE = RGBColor(0x2E, 0x75, 0xB6)         # bleu moyen
COULEUR_ENTETE_TABLEAU_HEX = "2E75B6"                   # même bleu, en hex
COULEUR_BANDEAU_HEX = "1F4E79"                          # bleu foncé, en hex

COULEUR_FOND_URGENCE_HEX = "FDEDE8"        # orange très clair
COULEUR_TEXTE_URGENCE = RGBColor(0xC0, 0x39, 0x2B)  # rouge/orange foncé

COULEUR_FOND_REGLE_OR_HEX = "EAF1FB"       # bleu très clair
COULEUR_TEXTE_REGLE_OR = COULEUR_TITRE_PRINCIPAL

COULEUR_FOND_PIEGE_HEX = "FFF6DA"          # jaune/ambre très clair
COULEUR_TEXTE_PIEGE = RGBColor(0x8A, 0x6D, 0x00)    # ambre foncé

COULEUR_LIGNE_ALTERNEE_HEX = "F2F6FC"      # gris-bleu très léger


def definir_couleur_fond_cellule(cellule, couleur_hex):

    tcPr = cellule._tc.get_or_add_tcPr()

    ombrage = OxmlElement("w:shd")

    ombrage.set(qn("w:val"), "clear")
    ombrage.set(qn("w:color"), "auto")
    ombrage.set(qn("w:fill"), couleur_hex)

    tcPr.append(ombrage)


def definir_marges_cellule(cellule, marge_dxa=150):

    tcPr = cellule._tc.get_or_add_tcPr()

    tcMar = OxmlElement("w:tcMar")

    for cote in ("top", "bottom", "left", "right"):

        element = OxmlElement(f"w:{cote}")

        element.set(qn("w:w"), str(marge_dxa))
        element.set(qn("w:type"), "dxa")

        tcMar.append(element)

    tcPr.append(tcMar)


TABLE_EXPOSANT = str.maketrans("0123456789+-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻")
TABLE_INDICE = str.maketrans("0123456789+-", "₀₁₂₃₄₅₆₇₈₉₊₋")


def _convertir_exposant(correspondance):

    contenu = correspondance.group(1)

    if re.fullmatch(r"[\d+\-]+", contenu):

        return contenu.translate(TABLE_EXPOSANT)

    return contenu


def _convertir_indice(correspondance):

    contenu = correspondance.group(1)

    if re.fullmatch(r"[\d+\-]+", contenu):

        return contenu.translate(TABLE_INDICE)

    return contenu


def nettoyer_latex(texte):

    remplacements = {
        r"\rightarrow": "→",
        r"\leftarrow": "←",
        r"\times": "×",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\approx": "≈",
        r"\pm": "±",
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\circ": "°",
    }

    for cle, valeur in remplacements.items():

        texte = texte.replace(cle, valeur)

    # \text{XYZ} → XYZ (retire l'enrobage, garde le contenu)
    texte = re.sub(r"\\text\{([^}]*)\}", r"\1", texte)

    # \, (espace fine LaTeX) → rien
    texte = texte.replace(r"\,", "")

    # Exposants/indices numériques : X^{14} → X¹⁴, X_2 → X₂
    # (vrais caractères exposant/indice, pas une simple
    # concaténation qui rendrait "10^{14}" en "1014").
    texte = re.sub(r"\^\{([^}]*)\}", _convertir_exposant, texte)
    texte = re.sub(r"_\{([^}]*)\}", _convertir_indice, texte)
    texte = re.sub(r"\^(\S)", _convertir_exposant, texte)
    texte = re.sub(r"_(\S)", _convertir_indice, texte)

    # Toute commande LaTeX restante non reconnue (\quelquechose)
    # est simplement retirée plutôt que laissée telle quelle.
    texte = re.sub(r"\\[a-zA-Z]+", "", texte)

    # Retire les délimiteurs $...$ ou $$...$$ sans supprimer
    # le contenu (on n'a pas de vrai rendu mathématique, mais
    # on évite au moins les "$" qui polluent le texte).
    texte = re.sub(r"\${1,2}", "", texte)

    # Nettoie les résidus isolés de ^ ou _ qui n'auraient pas
    # été capturés par les règles ci-dessus.
    texte = texte.replace("^", "").replace("_", "")

    return texte


MOTIF_BR = re.compile(r"<br\s*/?>", re.IGNORECASE)


def ajouter_run_avec_sauts(paragraphe, texte, gras=False, italique=False):

    # Convertit les <br> (ou <br/>, <br />) éventuellement
    # présents dans le texte — notamment dans les cellules de
    # tableau, où Gemini les utilise pour forcer un retour à
    # la ligne en Markdown — en véritables sauts de ligne
    # plutôt que de laisser la balise brute dans le document.

    parties = MOTIF_BR.split(texte)

    for i, partie in enumerate(parties):

        if partie:

            run = paragraphe.add_run(partie)

            run.bold = gras
            run.italic = italique

        if i < len(parties) - 1:

            paragraphe.add_run().add_break(WD_BREAK.LINE)


def ajouter_texte_formate(paragraphe, texte):

    # Découpe le texte en segments gras (**...**) / italique
    # (*...* ou _..._) / normal, et ajoute chaque segment
    # comme un "run" avec la bonne mise en forme.

    texte = nettoyer_latex(texte)

    motif = re.compile(
        r"(\*\*.+?\*\*|\*.+?\*|__.+?__|_.+?_)"
    )

    segments = motif.split(texte)

    for segment in segments:

        if not segment:

            continue

        if segment.startswith("**") and segment.endswith("**"):

            ajouter_run_avec_sauts(
                paragraphe, segment[2:-2], gras=True
            )

        elif segment.startswith("__") and segment.endswith("__"):

            ajouter_run_avec_sauts(
                paragraphe, segment[2:-2], gras=True
            )

        elif segment.startswith("*") and segment.endswith("*"):

            ajouter_run_avec_sauts(
                paragraphe, segment[1:-1], italique=True
            )

        elif segment.startswith("_") and segment.endswith("_"):

            ajouter_run_avec_sauts(
                paragraphe, segment[1:-1], italique=True
            )

        else:

            ajouter_run_avec_sauts(paragraphe, segment)


def est_ligne_tableau(ligne):

    return (
        ligne.strip().startswith("|")
        and ligne.strip().endswith("|")
    )


def est_ligne_separateur_tableau(ligne):

    contenu = ligne.strip().strip("|")

    cellules = contenu.split("|")

    return all(
        re.fullmatch(r"\s*:?-{2,}:?\s*", cellule)
        for cellule in cellules
    ) and len(cellules) > 0


def parser_ligne_tableau(ligne):

    contenu = ligne.strip().strip("|")

    return [
        cellule.strip()
        for cellule in contenu.split("|")
    ]


def creer_bandeau_titre(document, titre, sous_titre=None):

    table = document.add_table(rows=1, cols=1)

    cellule = table.cell(0, 0)

    definir_couleur_fond_cellule(cellule, COULEUR_BANDEAU_HEX)
    definir_marges_cellule(cellule, marge_dxa=250)

    p_titre = cellule.paragraphs[0]
    p_titre.alignment = 1  # centré

    run_titre = p_titre.add_run(titre)

    run_titre.bold = True
    run_titre.font.size = Pt(22)
    run_titre.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if sous_titre:

        p_sous_titre = cellule.add_paragraph()
        p_sous_titre.alignment = 1  # centré

        run_sous_titre = p_sous_titre.add_run(sous_titre)

        run_sous_titre.font.size = Pt(12)
        run_sous_titre.font.color.rgb = RGBColor(0xD9, 0xE7, 0xF5)


def creer_encadre(document, titre_encadre, icone,
                   couleur_fond_hex, couleur_texte):

    table = document.add_table(rows=1, cols=1)

    cellule = table.cell(0, 0)

    definir_couleur_fond_cellule(cellule, couleur_fond_hex)
    definir_marges_cellule(cellule, marge_dxa=150)

    p_titre = cellule.paragraphs[0]

    run_titre = p_titre.add_run(f"{icone} {titre_encadre}")

    run_titre.bold = True
    run_titre.font.size = Pt(13)
    run_titre.font.color.rgb = couleur_texte

    return cellule


def ajouter_tableau(document, lignes_tableau):

    entetes = parser_ligne_tableau(lignes_tableau[0])

    lignes_donnees = [
        parser_ligne_tableau(ligne)
        for ligne in lignes_tableau[2:]
    ]

    table = document.add_table(
        rows=1,
        cols=len(entetes)
    )

    table.style = "Table Grid"

    cellules_entete = table.rows[0].cells

    for i, texte_entete in enumerate(entetes):

        cellules_entete[i].paragraphs[0].text = ""

        ajouter_texte_formate(
            cellules_entete[i].paragraphs[0],
            texte_entete
        )

        for run in cellules_entete[i].paragraphs[0].runs:

            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        definir_couleur_fond_cellule(
            cellules_entete[i],
            COULEUR_ENTETE_TABLEAU_HEX
        )


    # Repère la colonne "gravité/complications" pour y ajouter
    # automatiquement une icône d'alerte.
    index_colonne_gravite = None

    for i, texte_entete in enumerate(entetes):

        if (
            "gravité" in texte_entete.lower()
            or "complication" in texte_entete.lower()
        ):

            index_colonne_gravite = i

            break


    for numero_ligne, ligne_donnees in enumerate(lignes_donnees):

        cellules = table.add_row().cells

        # Alternance de couleur (une ligne sur deux, en
        # commençant par une ligne colorée) pour la lisibilité.
        if numero_ligne % 2 == 0:

            for cellule in cellules:

                definir_couleur_fond_cellule(
                    cellule,
                    COULEUR_LIGNE_ALTERNEE_HEX
                )


        for i, valeur in enumerate(ligne_donnees):

            if i >= len(cellules):

                break

            if (
                i == index_colonne_gravite
                and valeur.strip()
                and "⚠️" not in valeur
            ):

                valeur = f"⚠️ {valeur}"

            cellules[i].paragraphs[0].text = ""

            ajouter_texte_formate(
                cellules[i].paragraphs[0],
                valeur
            )


def ajouter_sommaire(document, texte):

    titres_h1 = []

    for ligne in texte.split("\n"):

        ligne_strip = ligne.strip()

        if ligne_strip.startswith("#") and not ligne_strip.startswith("##"):

            titre = ligne_strip[1:].strip()

            if titre:

                titres_h1.append(titre)


    if not titres_h1:

        return


    titre_sommaire = document.add_heading("Sommaire", level=2)

    for run in titre_sommaire.runs:

        run.font.color.rgb = COULEUR_SOUS_TITRE


    for titre in titres_h1:

        p = document.add_paragraph(style="List Number")

        ajouter_texte_formate(p, titre)


def creer_word(texte, sous_titre=None):

    document = Document()

    creer_bandeau_titre(
        document,
        "Fiche de Révision IFSI",
        sous_titre
    )

    ajouter_sommaire(document, texte)

    lignes = texte.split("\n")

    dans_bloc_code = False
    tampon_code = []

    dans_tableau = False
    tampon_tableau = []

    # Conteneur où sont ajoutés les paragraphes courants :
    # le document lui-même, ou la cellule d'un encadré coloré
    # (Urgences / Règles d'or) le temps de cette section.
    conteneur_actif = document

    i = 0

    while i < len(lignes):

        ligne_brute = lignes[i]
        ligne = ligne_brute.rstrip()
        ligne_strip = ligne.strip()


        # ------------------------------------------------
        # BLOCS DE CODE ```...```
        # ------------------------------------------------

        if ligne_strip.startswith("```"):

            if dans_bloc_code:

                # Fin du bloc : on écrit le contenu en
                # police à chasse fixe.
                if tampon_code:

                    p = conteneur_actif.add_paragraph()

                    run = p.add_run(
                        "\n".join(tampon_code)
                    )

                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

                tampon_code = []
                dans_bloc_code = False

            else:

                dans_bloc_code = True

            i += 1

            continue

        if dans_bloc_code:

            tampon_code.append(ligne_brute)

            i += 1

            continue


        # ------------------------------------------------
        # TABLEAUX MARKDOWN
        # ------------------------------------------------

        if est_ligne_tableau(ligne_strip):

            if (
                not dans_tableau
                and i + 1 < len(lignes)
                and est_ligne_separateur_tableau(lignes[i + 1])
            ):

                dans_tableau = True
                tampon_tableau = [ligne_strip]

            elif dans_tableau:

                tampon_tableau.append(ligne_strip)

            i += 1

            continue

        elif dans_tableau:

            # Fin du tableau (ligne courante n'en fait
            # plus partie) : on le construit.
            ajouter_tableau(document, tampon_tableau)

            tampon_tableau = []
            dans_tableau = False

            # ne pas "continue" ici : on retraite la
            # ligne courante normalement ci-dessous


        if not ligne_strip:

            i += 1

            continue


        # ------------------------------------------------
        # LIGNES HORIZONTALES (---, ***, ___)
        # ------------------------------------------------

        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", ligne_strip):

            i += 1

            continue


        # ------------------------------------------------
        # TITRES MARKDOWN
        # ------------------------------------------------

        if ligne_strip.startswith("#"):

            niveau = len(
                ligne_strip
            ) - len(
                ligne_strip.lstrip("#")
            )

            niveau = max(1, min(niveau, 4))

            # Retire TOUS les groupes de # en tête de ligne,
            # même s'il y en a plusieurs séparés par des
            # espaces (ex : "## # Titre" → "Titre"), au cas
            # où le modèle produirait un formatage imparfait.
            titre_texte = re.sub(
                r"^(#+\s*)+",
                "",
                ligne_strip
            ).strip()

            # On revient au document normal avant de traiter
            # ce nouveau titre — un encadré éventuellement
            # ouvert par un titre précédent se referme ici.
            conteneur_actif = document

            titre_normalise = titre_texte.lower()

            if niveau >= 3 and "urgence" in titre_normalise and "vitale" in titre_normalise:

                conteneur_actif = creer_encadre(
                    document,
                    titre_texte,
                    "🚨",
                    COULEUR_FOND_URGENCE_HEX,
                    COULEUR_TEXTE_URGENCE
                )

                i += 1

                continue

            if niveau >= 3 and "règle" in titre_normalise and "or" in titre_normalise:

                conteneur_actif = creer_encadre(
                    document,
                    titre_texte,
                    "🔑",
                    COULEUR_FOND_REGLE_OR_HEX,
                    COULEUR_TEXTE_REGLE_OR
                )

                i += 1

                continue

            if niveau >= 3 and "piège" in titre_normalise:

                conteneur_actif = creer_encadre(
                    document,
                    titre_texte,
                    "⚡",
                    COULEUR_FOND_PIEGE_HEX,
                    COULEUR_TEXTE_PIEGE
                )

                i += 1

                continue

            p = document.add_heading(
                "",
                level=niveau
            )

            ajouter_texte_formate(p, titre_texte)

            couleur = (
                COULEUR_TITRE_PRINCIPAL
                if niveau == 1
                else COULEUR_SOUS_TITRE
            )

            for run in p.runs:

                run.font.color.rgb = couleur

            i += 1

            continue


        # ------------------------------------------------
        # LISTES À PUCES
        # ------------------------------------------------

        correspondance_puce = re.match(
            r"^[\-\*•]\s+(.*)",
            ligne_strip
        )

        if correspondance_puce:

            p = conteneur_actif.add_paragraph(
                style="List Bullet"
            )

            ajouter_texte_formate(
                p,
                correspondance_puce.group(1)
            )

            i += 1

            continue


        # ------------------------------------------------
        # LISTES NUMÉROTÉES
        # ------------------------------------------------

        correspondance_numero = re.match(
            r"^\d+[\.\)]\s+(.*)",
            ligne_strip
        )

        if correspondance_numero:

            p = conteneur_actif.add_paragraph(
                style="List Number"
            )

            ajouter_texte_formate(
                p,
                correspondance_numero.group(1)
            )

            i += 1

            continue


        # ------------------------------------------------
        # PARAGRAPHE NORMAL
        # ------------------------------------------------

        p = conteneur_actif.add_paragraph()

        ajouter_texte_formate(p, ligne_strip)

        i += 1


    # Si le texte se termine par un tableau non encore flushé
    if dans_tableau and tampon_tableau:

        ajouter_tableau(document, tampon_tableau)


    buffer = io.BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    return buffer


# ============================================================
# INTERFACE
# ============================================================

st.title(
    "📚 Générateur de Cours IFSI"
)


st.write(
    """
Colle l'URL d'une page de cours Lyon 1 (Moodle, "Livre", etc.)
— l'appli trouve automatiquement les vidéos qu'elle contient.
Tu peux aussi coller directement un ou plusieurs liens vidéo.
Une URL par ligne.
"""
)


urls_input = st.text_area(
    "URL de la page de cours (ou des vidéos)",
    height=200,
    placeholder=(
        "https://moodle.univ-lyon1.fr/mod/book/view.php?"
        "id=6424&chapterid=252\n"
        "https://...\n"
        "https://..."
    )
)


stockage_local = LocalStorage()

cle_gemini_memorisee = (
    stockage_local.getItem("ifsi_cle_gemini") or ""
)

cle_groq_memorisee = (
    stockage_local.getItem("ifsi_cle_groq") or ""
)

cle_api_utilisateur = st.text_input(
    "🔑 Ta clé API Gemini (gratuite, obligatoire)",
    type="password",
    value=cle_gemini_memorisee,
    placeholder="ex : AIzaSyD-9xY2kLmN3pQ4rS5tU6vW7xY8zA9bC0"
)

cle_api_groq = st.text_input(
    "🔑 Ta clé API Groq (gratuite, optionnelle — rend la "
    "transcription beaucoup plus rapide)",
    type="password",
    value=cle_groq_memorisee,
    placeholder="ex : gsk_..."
)

memoriser_cles = st.checkbox(
    "💾 Se souvenir de mes clés sur cet appareil (pour ne "
    "pas avoir à les recoller à chaque fois)",
    value=bool(cle_gemini_memorisee or cle_groq_memorisee)
)

st.caption(
    "⚠️ Ne coche pas cette case sur un ordinateur partagé "
    "(bibliothèque, salle informatique) — tes clés seraient "
    "visibles pour la prochaine personne qui utilise ce "
    "navigateur sur cette machine."
)

if memoriser_cles:

    if (
        cle_api_utilisateur
        and cle_api_utilisateur != cle_gemini_memorisee
    ):

        stockage_local.setItem(
            "ifsi_cle_gemini",
            cle_api_utilisateur,
            key="enregistrer_cle_gemini"
        )

    if (
        cle_api_groq
        and cle_api_groq != cle_groq_memorisee
    ):

        stockage_local.setItem(
            "ifsi_cle_groq",
            cle_api_groq,
            key="enregistrer_cle_groq"
        )

elif cle_gemini_memorisee or cle_groq_memorisee:

    # La case a été décochée alors que des clés étaient
    # enregistrées : on les efface du navigateur.
    stockage_local.deleteItem(
        "ifsi_cle_gemini",
        key="oublier_cle_gemini"
    )

    stockage_local.deleteItem(
        "ifsi_cle_groq",
        key="oublier_cle_groq"
    )

with st.expander("ℹ️ Comment obtenir mes clés API ?"):

    st.write(
        "**Clé Gemini (obligatoire)** — obtiens-la "
        "gratuitement sur https://aistudio.google.com/apikey "
        "(aucune carte bancaire requise). Elle n'est jamais "
        "enregistrée par cette appli.\n\n"
        "**Clé Groq (optionnelle)** — obtiens-la gratuitement "
        "sur https://console.groq.com/keys (aucune carte "
        "bancaire requise). Elle sert uniquement à accélérer "
        "la transcription audio (quelques secondes au lieu "
        "de plusieurs minutes). **Si tu laisses ce champ "
        "vide, l'appli fonctionne quand même** — juste plus "
        "lentement."
    )


MODELES_DISPONIBLES = {
    "gemini-3.8-flash": (
        "Gemini 3.8 Flash — recommandé (le plus récent)"
    ),
    "gemini-3.7-flash": (
        "Gemini 3.7 Flash — éprouvé, très fiable"
    ),
    "gemini-3.6-flash": (
        "Gemini 3.6 Flash — génération précédente, bon repli"
    ),
    "gemini-3.5-flash-lite": (
        "Gemini 3.5 Flash-Lite — le plus léger"
    ),
}

# Chaîne de repli ORDONNÉE, utilisée quand le modèle choisi
# devient indisponible (503 persistant) ou a épuisé son
# quota journalier (429). Chaque modèle a son propre quota
# séparé sur le niveau gratuit, donc basculer vers le suivant
# de la chaîne peut débloquer la situation sans attendre ni
# payer.
#
# Note : gemini-2.5-flash a été retiré de cette liste — Google
# a fermé ce modèle aux nouveaux comptes (404 NOT_FOUND,
# message recommandant gemini-3.6-flash à la place). Toujours
# vérifier qu'un modèle existe encore avant de l'ajouter à
# cette chaîne : les identifiants de modèles Gemini évoluent.
CHAINE_MODELES_SECOURS = [
    "gemini-3.8-flash",
    "gemini-3.7-flash",
    "gemini-3.6-flash",
    "gemini-3.5-flash-lite",
]


def modele_secours_suivant(modele_actuel):

    if modele_actuel not in CHAINE_MODELES_SECOURS:

        return None

    index = CHAINE_MODELES_SECOURS.index(modele_actuel)

    if index + 1 < len(CHAINE_MODELES_SECOURS):

        return CHAINE_MODELES_SECOURS[index + 1]

    return None


MODES_TRAITEMENT = {
    "rapide": "⚡ Rapide (audio envoyé à Gemini)",
    "gratuit": "🆓 Gratuit (Groq + repli local si besoin)",
}

# Réglages avancés — repliés par défaut : la plupart des
# étudiants n'ont pas besoin d'y toucher, les valeurs par
# défaut (Gemini 3.8 Flash, mode gratuit) conviennent pour
# l'immense majorité des cas.
with st.expander("⚙️ Réglages avancés (facultatif)"):

    modele_choisi = st.selectbox(
        "🤖 Modèle Gemini",
        options=list(MODELES_DISPONIBLES.keys()),
        format_func=lambda cle: MODELES_DISPONIBLES[cle],
        index=0
    )

    st.caption(
        "D'après nos tests sur des cours à plusieurs sujets : "
        "**Gemini 3.8 Flash** est le plus récent (recommandé) "
        ", **3.7 Flash** est celui qu'on a le plus testé "
        "(très fiable), **3.6 Flash** est un bon repli si les "
        "derniers modèles sont saturés, et **3.5 Flash-Lite** "
        "est le plus léger mais a parfois oublié des sujets "
        "sur des cours couvrant plusieurs thèmes — à réserver "
        "aux cours courts ou en dernier recours."
    )

    mode_traitement = st.radio(
        "Mode de traitement de l'audio",
        options=list(MODES_TRAITEMENT.keys()),
        format_func=lambda cle: MODES_TRAITEMENT[cle],
        index=1,
        horizontal=True
    )

    st.caption(
        "**⚡ Rapide** : l'audio est envoyé directement à "
        "Gemini — quelques minutes de traitement, mais "
        "consomme davantage de quota Gemini (plus susceptible "
        "de tomber sur un quota épuisé ou une surcharge). "
        "**🆓 Gratuit** (par défaut) : l'audio est transcrit "
        "via Groq si une clé est fournie (rapide), sinon en "
        "local sur ce serveur (plus lent) — Gemini ne reçoit "
        "ensuite que du texte, donc bien moins de risque de "
        "saturation."
    )


# Note : pas d'authentification nécessaire — les liens MP3
# publics et l'API "modes" (URLs signées) fonctionnent sans
# cookie de session sur cette plateforme.


# ============================================================
# FONCTION : COLLECTE + EMPAQUETAGE "SANS IA"
# ============================================================
#
# Chemin alternatif proposé dès le départ (pas seulement en
# secours après un échec Gemini) : transcrit l'audio (Groq
# prioritaire, repli local) et récupère les PDF de support,
# sans jamais appeler Gemini. Retourne un .zip contenant le
# prompt complet (.txt, prêt à coller dans une autre IA) et
# les PDF originaux. Utile quand Gemini est saturé et que
# l'utilisateur ne veut pas perdre de temps à attendre.
# ============================================================

# ============================================================
# CACHE PARTAGÉ ENTRE LES DEUX BOUTONS ("avec IA" / "sans IA")
# ============================================================
#
# Évite de re-télécharger/re-transcrire si l'un des deux
# chemins a déjà collecté les transcriptions + PDF pour les
# mêmes URL (typiquement : Gemini a été tenté et a échoué en
# mode gratuit, puis l'utilisateur clique sur "sans IA" — ou
# l'inverse). Uniquement valable pour le contenu textuel du
# mode gratuit (transcriptions + texte PDF) : en mode rapide,
# les contenus sont des fichiers déjà envoyés à Gemini, donc
# rien de commun à réutiliser ici.

def cle_cache_collecte(urls_brutes):

    return "\n".join(urls_brutes)


def lire_cache_collecte(urls_brutes):

    cache = st.session_state.get("cache_collecte_sans_ia")

    if cache and cache.get("cle") == cle_cache_collecte(
        urls_brutes
    ):

        return cache

    return None


def ecrire_cache_collecte(
    urls_brutes,
    contenus_textuels,
    pdfs_bruts,
    titre_cours
):

    st.session_state["cache_collecte_sans_ia"] = {
        "cle": cle_cache_collecte(urls_brutes),
        "contenus_textuels": contenus_textuels,
        "pdfs_bruts": pdfs_bruts,
        "titre_cours": titre_cours,
    }


def collecter_et_empaqueter_sans_ia(urls_brutes, cle_api_groq):

    cache = lire_cache_collecte(urls_brutes)

    if cache:

        st.info(
            "♻️ Réutilisation des transcriptions et PDF déjà "
            "collectés pour ces URL — aucun nouveau "
            "téléchargement ni appel réseau."
        )

        contenus_textuels = cache["contenus_textuels"]
        pdfs_bruts = cache["pdfs_bruts"]
        titre_cours = cache["titre_cours"]

        return empaqueter_zip(
            contenus_textuels,
            pdfs_bruts,
            titre_cours
        )


    session = requests.Session()

    session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 "
            "(KHTML, like Gecko) "
            "Chrome/140.0 Safari/537.36"
        )
    })

    urls, pdfs_detectes, titre_cours = developper_urls(
        urls_brutes,
        session
    )

    if not urls:

        st.error(
            "❌ Aucun lien vidéo n'a pu être trouvé à partir "
            "des URL fournies."
        )

        return None, None


    contenus_textuels = []
    pdfs_bruts = []  # liste de (nom_fichier, octets)

    with st.status(
        "Collecte des documents (sans IA)...",
        expanded=True
    ) as status:

        for i, url in enumerate(urls):

            st.write(
                f"## 🎧 Cours {i + 1}/{len(urls)}"
            )

            fichier_local = recuperer_audio(
                url,
                i,
                session
            )

            if not fichier_local:

                st.warning(
                    f"⚠️ Aucun audio trouvé pour : {url}"
                )

                continue

            if ACCELERATION_ACTIVEE:

                fichier_local = accelerer_audio(
                    fichier_local,
                    i + 1
                )

            texte_audio = transcrire_audio(
                fichier_local,
                i + 1,
                cle_api_groq
            )

            try:

                os.remove(fichier_local)

            except Exception:

                pass

            if texte_audio:

                contenus_textuels.append(
                    f"==============================\n"
                    f"TRANSCRIPTION AUDIO {i + 1}\n"
                    f"==============================\n\n"
                    f"{texte_audio}"
                )


        if pdfs_detectes:

            st.write(
                f"## 📄 Supports PDF ({len(pdfs_detectes)})"
            )

            for j, url_pdf in enumerate(pdfs_detectes):

                texte_pdf, octets_pdf = (
                    telecharger_et_extraire_pdf(
                        url_pdf,
                        j + 1,
                        session,
                        retourner_bytes=True
                    )
                )

                if texte_pdf:

                    contenus_textuels.append(
                        f"==========================\n"
                        f"SUPPORT PDF {j + 1}\n"
                        f"==========================\n\n"
                        f"{texte_pdf}"
                    )

                if octets_pdf:

                    pdfs_bruts.append(
                        (f"support_{j + 1}.pdf", octets_pdf)
                    )


        if not contenus_textuels and not pdfs_bruts:

            status.update(
                label="❌ Rien n'a pu être récupéré",
                state="error",
                expanded=True
            )

            st.error(
                "Aucune transcription ni aucun support PDF "
                "n'a pu être obtenu."
            )

            return None, None

        status.update(
            label="✅ Documents collectés avec succès !",
            state="complete"
        )


    ecrire_cache_collecte(
        urls_brutes,
        contenus_textuels,
        pdfs_bruts,
        titre_cours
    )


    return empaqueter_zip(
        contenus_textuels,
        pdfs_bruts,
        titre_cours
    )


def empaqueter_zip(contenus_textuels, pdfs_bruts, titre_cours):

    tampon_zip = io.BytesIO()

    with zipfile.ZipFile(
        tampon_zip,
        "w",
        zipfile.ZIP_DEFLATED
    ) as archive:

        if contenus_textuels:

            prompt_txt = construire_prompt_fiche_texte(
                contenus_textuels
            )

            archive.writestr(
                "transcriptions_cours_ifsi.txt",
                prompt_txt
            )

        for nom_fichier, octets_pdf in pdfs_bruts:

            archive.writestr(
                nom_fichier,
                octets_pdf
            )

    tampon_zip.seek(0)

    return tampon_zip.getvalue(), titre_cours


# ============================================================
# BOUTONS
# ============================================================

colonne_ia, colonne_sans_ia = st.columns(2)

with colonne_sans_ia:

    telechargement_demande = st.button(
        "📥 Télécharger sans passer par l'IA",
        help=(
            "Récupère directement les transcriptions et les "
            "PDF de support dans un .zip, sans appeler Gemini "
            "— utile si Gemini est saturé ou pour aller plus "
            "vite."
        )
    )

if telechargement_demande:

    urls_zip = [
        url.strip()
        for url in urls_input.splitlines()
        if url.strip()
    ]

    if not urls_zip:

        st.warning(
            "⚠️ Ajoute au moins une URL."
        )

        st.stop()

    contenu_zip, titre_cours_zip = (
        collecter_et_empaqueter_sans_ia(
            urls_zip,
            cle_api_groq
        )
    )

    if contenu_zip:

        st.download_button(
            label="📦 Télécharger le .zip (transcriptions + PDF)",
            data=contenu_zip,
            file_name=(
                nettoyer_nom_fichier(titre_cours_zip)
                + "_documents.zip"
            ),
            mime="application/zip"
        )


with colonne_ia:

    bouton_generer = st.button(
        "🚀 Générer la fiche complète",
        type="primary"
    )

if bouton_generer:


    # --------------------------------------------------------
    # URL
    # --------------------------------------------------------

    urls = [
        url.strip()
        for url in urls_input.splitlines()
        if url.strip()
    ]


    if not urls:

        st.warning(
            "⚠️ Ajoute au moins une URL."
        )

        st.stop()

    # Conservé tel quel (avant expansion par developper_urls)
    # pour servir de clé de cache commune avec le bouton
    # "sans IA", qui utilise les mêmes URL brutes.
    urls_brutes_saisies = urls


    # --------------------------------------------------------
    # SESSION HTTP (créée tôt : nécessaire pour développer
    # les pages de cours en liens vidéo, avant même de créer
    # le client Gemini)
    # --------------------------------------------------------

    session = requests.Session()

    session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 "
            "(KHTML, like Gecko) "
            "Chrome/140.0 Safari/537.36"
        )
    })


    # --------------------------------------------------------
    # DÉVELOPPEMENT DES URLS (pages de cours → liens vidéo)
    # --------------------------------------------------------

    urls, pdfs_detectes, titre_cours = developper_urls(
        urls,
        session
    )

    if not urls:

        st.error(
            "❌ Aucun lien vidéo n'a pu être trouvé à partir "
            "des URL fournies."
        )

        st.stop()


    # --------------------------------------------------------
    # API KEY
    # --------------------------------------------------------

    api_key = (
        cle_api_utilisateur.strip()
        if cle_api_utilisateur
        else None
    )

    if not api_key:

        try:

            api_key = st.secrets[
                "GEMINI_API_KEY"
            ]

        except Exception:

            api_key = None

    if not api_key:

        st.error(
            """
❌ Aucune clé API Gemini renseignée.

Colle ta clé dans le champ "🔑 Ta clé API Gemini" ci-dessus.

Tu peux en obtenir une gratuitement (sans carte bancaire) sur :

https://aistudio.google.com/apikey
"""
        )

        st.stop()


    # --------------------------------------------------------
    # CLIENT
    # --------------------------------------------------------

    try:

        client = genai.Client(
            api_key=api_key
        )

    except Exception as e:

        st.error(
            f"❌ Impossible de créer le client Gemini : {e}"
        )

        st.stop()


    # Liste des contenus à analyser : soit des blocs de texte
    # (mode gratuit — transcriptions + PDF extraits en local),
    # soit des fichiers Gemini (mode rapide — audio/PDF envoyés
    # directement), selon le mode choisi par l'utilisateur.
    contenus_a_traiter = []

    # Octets bruts des PDF de support (mode gratuit uniquement)
    # — alimente le cache partagé avec le bouton "sans IA".
    pdfs_bruts_gratuit = []

    # Fichiers Gemini des PDF de support (mode gratuit) — les
    # PDF sont désormais envoyés à Gemini en fichiers natifs
    # (schémas/images inclus), pas seulement en texte extrait.
    # Attachés au dernier appel Gemini par generer_fiche_finale.
    fichiers_pdf_gemini_gratuit = []


    # ========================================================
    # TRAITEMENT
    # ========================================================

    activer_autoscroll()

    debut_total = time.time()

    with st.status(
        "Traitement des cours...",
        expanded=True
    ) as status:


        # ----------------------------------------------------
        # POUR CHAQUE URL : récupération + upload seulement
        # (l'analyse se fait plus bas, en un seul appel)
        # ----------------------------------------------------

        for i, url in enumerate(urls):


            st.write(
                f"## 🎧 Cours {i + 1}/{len(urls)}"
            )


            # ----------------------------------------------
            # Téléchargement
            # ----------------------------------------------

            debut_audio = time.time()

            fichier_local = recuperer_audio(
                url,
                i,
                session
            )

            duree_audio = time.time() - debut_audio


            if not fichier_local:

                st.warning(
                    f"⚠️ Aucun audio récupérable (ni MP3 direct, "
                    f"ni HLS) pour : {url} "
                    f"(⏱️ {duree_audio:.0f}s)"
                )

                continue

            st.write(
                f"⏱️ Récupération de l'audio : {duree_audio:.0f}s"
            )


            # ----------------------------------------------
            # Accélération (accélère aussi la transcription
            # locale, puisqu'il y a moins de secondes à traiter)
            # ----------------------------------------------

            if ACCELERATION_ACTIVEE:

                fichier_local = accelerer_audio(
                    fichier_local,
                    i + 1
                )


            if mode_traitement == "gratuit":

                # ------------------------------------------
                # Transcription : Groq en priorité (rapide,
                # gratuit, hors de notre hébergement), repli
                # automatique sur Whisper local si besoin.
                # ------------------------------------------

                debut_transcription = time.time()

                texte_audio = transcrire_audio(
                    fichier_local,
                    i + 1,
                    cle_api_groq
                )

                duree_transcription = (
                    time.time() - debut_transcription
                )

                st.write(
                    f"⏱️ Transcription de l'audio {i + 1} : "
                    f"{duree_transcription:.0f}s"
                )

                try:

                    os.remove(fichier_local)

                except Exception:

                    pass

                if texte_audio:

                    contenus_a_traiter.append(
                        f"==============================\n"
                        f"TRANSCRIPTION AUDIO {i + 1}\n"
                        f"==============================\n\n"
                        f"{texte_audio}"
                    )

            else:

                # ------------------------------------------
                # Mode rapide : envoi direct à Gemini
                # ------------------------------------------

                debut_upload = time.time()

                fichier_gemini = uploader_audio_gemini(
                    client,
                    fichier_local,
                    i + 1
                )

                duree_upload = time.time() - debut_upload

                st.write(
                    f"⏱️ Upload de l'audio {i + 1} : "
                    f"{duree_upload:.0f}s"
                )

                try:

                    os.remove(fichier_local)

                except Exception:

                    pass

                if fichier_gemini:

                    contenus_a_traiter.append(fichier_gemini)


        # ----------------------------------------------------
        # PDF SUPPORT (diapositives, si trouvées sur la page)
        # ----------------------------------------------------

        if pdfs_detectes:

            st.write(
                f"## 📄 Supports PDF ({len(pdfs_detectes)})"
            )

            for j, url_pdf in enumerate(pdfs_detectes):

                if mode_traitement == "gratuit":

                    # PDF envoyé à Gemini en fichier natif
                    # (schémas/images inclus) — un seul appel
                    # generate_content au total malgré tout,
                    # l'upload de fichier n'y est pas soumis.
                    fichier_pdf_gemini, octets_pdf = (
                        telecharger_et_uploader_pdf(
                            url_pdf,
                            j + 1,
                            session,
                            client,
                            retourner_bytes=True
                        )
                    )

                    if fichier_pdf_gemini:

                        fichiers_pdf_gemini_gratuit.append(
                            fichier_pdf_gemini
                        )

                    if octets_pdf:

                        pdfs_bruts_gratuit.append(
                            (f"support_{j + 1}.pdf", octets_pdf)
                        )

                else:

                    fichier_pdf_gemini = (
                        telecharger_et_uploader_pdf(
                            url_pdf,
                            j + 1,
                            session,
                            client
                        )
                    )

                    if fichier_pdf_gemini:

                        contenus_a_traiter.append(
                            fichier_pdf_gemini
                        )


        # ----------------------------------------------------
        # Alimente le cache partagé avec le bouton "sans IA"
        # (uniquement pertinent en mode gratuit : contenus déjà
        # textuels, réutilisables tels quels pour le .txt/.zip)
        # ----------------------------------------------------

        if mode_traitement == "gratuit" and (
            contenus_a_traiter or pdfs_bruts_gratuit
        ):

            ecrire_cache_collecte(
                urls_brutes_saisies,
                contenus_a_traiter,
                pdfs_bruts_gratuit,
                titre_cours
            )


        # ====================================================
        # VÉRIFICATION
        # ====================================================

        if not contenus_a_traiter and not fichiers_pdf_gemini_gratuit:

            status.update(
                label="❌ Aucun cours analysé",
                state="error",
                expanded=True
            )

            st.error(
                """
Aucune transcription ni aucun support PDF n'a pu être
obtenu.

Vérifie notamment que les URL contiennent bien
un lien vers un fichier MP3 accessible.
"""
            )

            st.stop()


        # ====================================================
        # FICHE FINALE (un seul appel Gemini)
        # ====================================================

        debut_fiche = time.time()

        try:

            fiche = generer_fiche_finale(
                client,
                contenus_a_traiter,
                modele_choisi,
                mode_traitement,
                fichiers_pdf_gemini_gratuit
            )

        except (
            QuotaEpuiseeError,
            genai_errors.ServerError,
            ModeleIndisponibleError
        ) as e:

            # On garde les contenus déjà obtenus (transcriptions
            # ou fichiers Gemini déjà uploadés) pour permettre
            # un nouvel essai avec un autre modèle sans tout
            # refaire.
            st.session_state["contenus_textuels_en_attente"] = (
                contenus_a_traiter
            )

            st.session_state["pdfs_bruts_en_attente"] = (
                pdfs_bruts_gratuit
            )

            st.session_state["fichiers_pdf_gemini_en_attente"] = (
                fichiers_pdf_gemini_gratuit
            )

            st.session_state["mode_en_attente"] = mode_traitement

            st.session_state["modele_echoue"] = modele_choisi

            st.session_state["titre_cours_en_attente"] = (
                titre_cours
            )

            if isinstance(e, QuotaEpuiseeError):

                st.session_state["type_erreur_gemini"] = "quota"

            elif isinstance(e, ModeleIndisponibleError):

                st.session_state["type_erreur_gemini"] = (
                    "modele_indisponible"
                )

            else:

                st.session_state["type_erreur_gemini"] = (
                    "surcharge"
                )

            # st.rerun() (pas st.stop()) : on a besoin que le
            # script redémarre pour atteindre la section plus
            # bas qui affiche le message d'erreur et le bouton
            # "Réessayer avec..." — st.stop() figerait la page
            # ici et cette section n'apparaîtrait jamais.
            st.rerun()

        duree_fiche = time.time() - debut_fiche

        st.write(
            f"⏱️ Analyse + création de la fiche : "
            f"{duree_fiche:.0f}s"
        )


        if not fiche:

            # Même filet de sécurité que pour les erreurs
            # anticipées : on garde les contenus déjà obtenus
            # pour permettre un nouvel essai sans tout refaire,
            # même si l'échec vient d'une cause imprévue (pas
            # juste quota/surcharge/modèle indisponible).
            st.session_state["contenus_textuels_en_attente"] = (
                contenus_a_traiter
            )

            st.session_state["pdfs_bruts_en_attente"] = (
                pdfs_bruts_gratuit
            )

            st.session_state["fichiers_pdf_gemini_en_attente"] = (
                fichiers_pdf_gemini_gratuit
            )

            st.session_state["mode_en_attente"] = mode_traitement

            st.session_state["modele_echoue"] = modele_choisi

            st.session_state["titre_cours_en_attente"] = (
                titre_cours
            )

            st.session_state["type_erreur_gemini"] = "generique"

            # st.rerun() (pas st.stop()) : on a besoin que le
            # script redémarre pour atteindre la section plus
            # bas qui affiche le message d'erreur et le bouton
            # de reprise.
            st.rerun()


        # Nettoyage des fichiers Gemini : en mode rapide, tout
        # contenus_a_traiter ; en mode gratuit, uniquement les
        # PDF (les transcriptions restent de simples chaînes,
        # rien à supprimer côté Gemini pour elles).
        if mode_traitement == "rapide":

            elements_a_nettoyer = contenus_a_traiter

        else:

            elements_a_nettoyer = fichiers_pdf_gemini_gratuit

        for element in elements_a_nettoyer:

            try:

                client.files.delete(name=element.name)

            except Exception:

                pass


        duree_totale = time.time() - debut_total

        st.write(
            f"⏱️ **Durée totale du traitement : "
            f"{duree_totale / 60:.1f} min "
            f"({duree_totale:.0f}s)**"
        )

        status.update(
            label=(
                f"✅ Fiche générée avec succès en "
                f"{duree_totale / 60:.1f} min !"
            ),
            state="complete",
            expanded=True
        )


    # ========================================================
    # SAUVEGARDE EN SESSION (survit au clic sur "Télécharger")
    # ========================================================

    st.session_state["fiche_generee"] = fiche

    st.session_state["nom_fichier_docx"] = (
        nettoyer_nom_fichier(titre_cours) + ".docx"
    )


# ============================================================
# REPRISE AVEC MODÈLE DE SECOURS (hors du bloc bouton, pour
# rester actionnable après un échec quota/surcharge)
# ============================================================

if st.session_state.get("contenus_textuels_en_attente"):

    type_erreur = st.session_state.get("type_erreur_gemini")

    if type_erreur == "quota":

        afficher_erreur_quota()

    elif type_erreur == "modele_indisponible":

        afficher_erreur_modele_indisponible()

    elif type_erreur == "surcharge":

        afficher_erreur_surcharge()

    else:

        st.error(
            "❌ La création de la fiche a échoué pour une "
            "raison imprévue. Tes transcriptions sont "
            "conservées."
        )

    mode_en_attente_export = st.session_state.get(
        "mode_en_attente", "gratuit"
    )

    # Filet de sécurité indépendant de Gemini : uniquement en
    # mode 🆓 gratuit, les contenus en attente sont du texte
    # brut (transcriptions + PDF), pas des fichiers déjà
    # envoyés à Gemini — donc exportables tels quels. En mode
    # ⚡ rapide, les contenus sont des références de fichiers
    # Gemini déjà uploadés : rien de local à proposer ici.
    # Même .zip (transcriptions .txt + PDF originaux) que le
    # bouton "sans IA", pour que l'utilisateur retrouve
    # exactement les mêmes fichiers, qu'il ait cliqué sur
    # "sans IA" directement ou qu'il soit tombé sur le mur de
    # la surcharge Gemini.
    if mode_en_attente_export == "gratuit":

        zip_a_exporter, _ = empaqueter_zip(
            st.session_state["contenus_textuels_en_attente"],
            st.session_state.get("pdfs_bruts_en_attente", []),
            st.session_state.get(
                "titre_cours_en_attente", "Cours"
            )
        )

        st.download_button(
            label=(
                "📥 Télécharger les documents "
                "(transcriptions + PDF)"
            ),
            data=zip_a_exporter,
            file_name=(
                nettoyer_nom_fichier(
                    st.session_state.get(
                        "titre_cours_en_attente", "Cours"
                    )
                ) + "_documents.zip"
            ),
            mime="application/zip",
            help=(
                "Contient les transcriptions, les supports PDF "
                "et les consignes déjà rédigées pour la fiche — "
                "colle le .txt tel quel dans ChatGPT, Claude, "
                "Gemini web, etc."
            )
        )

    modele_echoue = st.session_state.get("modele_echoue")
    modele_secours = modele_secours_suivant(modele_echoue)

    if modele_secours:

        libelle_secours = MODELES_DISPONIBLES.get(
            modele_secours, modele_secours
        )

        if st.button(
            f"🔄 Réessayer avec {libelle_secours}"
        ):

            cle_api_reprise = (
                cle_api_utilisateur.strip()
                if cle_api_utilisateur
                else None
            )

            if not cle_api_reprise:

                try:

                    cle_api_reprise = st.secrets[
                        "GEMINI_API_KEY"
                    ]

                except Exception:

                    cle_api_reprise = None

            if not cle_api_reprise:

                st.error(
                    "❌ Clé API introuvable pour réessayer — "
                    "recolle-la dans le champ ci-dessus."
                )

            else:

                contenus_en_attente = st.session_state[
                    "contenus_textuels_en_attente"
                ]

                pdf_gemini_en_attente = st.session_state.get(
                    "fichiers_pdf_gemini_en_attente", []
                )

                mode_en_attente = st.session_state.get(
                    "mode_en_attente", "gratuit"
                )

                client_reprise = genai.Client(
                    api_key=cle_api_reprise
                )

                fiche_reprise = None
                echec_reprise = False

                with st.spinner(
                    f"Nouvelle tentative avec "
                    f"{libelle_secours}..."
                ):

                    try:

                        fiche_reprise = generer_fiche_finale(
                            client_reprise,
                            contenus_en_attente,
                            modele_secours,
                            mode_en_attente,
                            pdf_gemini_en_attente
                        )

                    except QuotaEpuiseeError:

                        echec_reprise = True

                        st.session_state["type_erreur_gemini"] = (
                            "quota"
                        )

                    except ModeleIndisponibleError:

                        echec_reprise = True

                        st.session_state["type_erreur_gemini"] = (
                            "modele_indisponible"
                        )

                    except genai_errors.ServerError:

                        echec_reprise = True

                        st.session_state["type_erreur_gemini"] = (
                            "surcharge"
                        )


                if fiche_reprise:

                    # Succès : nettoyage des fichiers Gemini —
                    # tout contenus_en_attente en mode rapide,
                    # seulement les PDF en mode gratuit.
                    if mode_en_attente == "rapide":

                        elements_a_nettoyer = contenus_en_attente

                    else:

                        elements_a_nettoyer = (
                            pdf_gemini_en_attente
                        )

                    for element in elements_a_nettoyer:

                        try:

                            client_reprise.files.delete(
                                name=element.name
                            )

                        except Exception:

                            pass

                    del st.session_state[
                        "contenus_textuels_en_attente"
                    ]

                    del st.session_state["modele_echoue"]

                    st.session_state.pop(
                        "type_erreur_gemini", None
                    )

                    st.session_state.pop(
                        "mode_en_attente", None
                    )

                    st.session_state.pop(
                        "pdfs_bruts_en_attente", None
                    )

                    st.session_state.pop(
                        "fichiers_pdf_gemini_en_attente", None
                    )

                    titre_cours_attente = st.session_state.pop(
                        "titre_cours_en_attente",
                        None
                    )

                    st.session_state["fiche_generee"] = (
                        fiche_reprise
                    )

                    st.session_state["nom_fichier_docx"] = (
                        nettoyer_nom_fichier(
                            titre_cours_attente
                        ) + ".docx"
                    )

                    st.rerun()

                else:

                    # Échec — connu (quota/surcharge/modèle
                    # indisponible) ou imprévu, peu importe :
                    # on ne perd jamais les transcriptions
                    # tant qu'il reste un modèle dans la chaîne.
                    if not echec_reprise:

                        st.session_state["type_erreur_gemini"] = (
                            "generique"
                        )

                    if modele_secours_suivant(modele_secours):

                        # Il reste un modèle dans la chaîne :
                        # on garde les contenus et on avance
                        # d'un cran pour proposer le suivant.
                        st.session_state["modele_echoue"] = (
                            modele_secours
                        )

                        st.rerun()

                    else:

                        # Fin de la chaîne : on abandonne, et
                        # on nettoie les fichiers Gemini restants
                        # (tout en mode rapide, PDF seulement en
                        # mode gratuit).
                        if mode_en_attente == "rapide":

                            elements_a_nettoyer = (
                                contenus_en_attente
                            )

                        else:

                            elements_a_nettoyer = (
                                pdf_gemini_en_attente
                            )

                        for element in elements_a_nettoyer:

                            try:

                                client_reprise.files.delete(
                                    name=element.name
                                )

                            except Exception:

                                pass

                        del st.session_state[
                            "contenus_textuels_en_attente"
                        ]

                        del st.session_state["modele_echoue"]

                        st.session_state.pop(
                            "type_erreur_gemini", None
                        )

                        st.session_state.pop(
                            "mode_en_attente", None
                        )

                        st.session_state.pop(
                            "pdfs_bruts_en_attente", None
                        )

                        st.session_state.pop(
                            "fichiers_pdf_gemini_en_attente",
                            None
                        )

                        st.session_state.pop(
                            "titre_cours_en_attente", None
                        )

                        st.error(
                            "❌ Tous les modèles de la chaîne "
                            "de repli sont actuellement "
                            "indisponibles."
                        )


# ============================================================
# AFFICHAGE DU RÉSULTAT (hors du bloc bouton, pour survivre
# au clic sur "Télécharger" qui relance le script)
# ============================================================

if st.session_state.get("fiche_generee"):

    st.success(
        "🎉 Ta fiche de révision est prête !"
    )

    st.markdown(
        st.session_state["fiche_generee"]
    )


    st.write(
        "### 📄 Télécharger la fiche"
    )

    nom_sans_extension = st.session_state["nom_fichier_docx"]

    if nom_sans_extension.lower().endswith(".docx"):

        nom_sans_extension = nom_sans_extension[:-5]

    buffer = creer_word(
        st.session_state["fiche_generee"],
        sous_titre=nom_sans_extension
    )

    colonne_telecharger, colonne_recharger = st.columns(2)

    with colonne_telecharger:

        st.download_button(
            label="📥 Télécharger la fiche (.docx)",
            data=buffer,
            file_name=st.session_state["nom_fichier_docx"],
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            )
        )

    with colonne_recharger:

        if st.button("🔄 Nouveau cours"):

            st.session_state.pop("fiche_generee", None)
            st.session_state.pop("nom_fichier_docx", None)

            st.rerun()